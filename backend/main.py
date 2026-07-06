import os
import re
import json
import random
import asyncio
import datetime
import ssl
import nltk
from nltk.corpus import wordnet as wn

try:
    nltk.data.find('corpora/wordnet')
except LookupError:
    ssl._create_default_https_context = ssl._create_unverified_context
    nltk.download('wordnet', quiet=True)
    nltk.download('omw-1.4', quiet=True)

import difflib
try:
    ALL_LEMMAS = list(set(wn.all_lemma_names()))
except Exception:
    ALL_LEMMAS = []

from fastapi import FastAPI, HTTPException, Header, UploadFile, File, Depends
from pydantic import BaseModel, Field
from dotenv import load_dotenv

from google.antigravity import Agent, LocalAgentConfig, CapabilitiesConfig

# Single-shot generation config: disable the agent's builtin tools/subagents so we get a
# fast single model response (~8s) instead of a full coding-agent loop (~40s).
FAST_CAPS = CapabilitiesConfig(enable_subagents=False, enabled_tools=[])


# AI-token accounting: a small JSON file next to the backend (no DB schema change).
# Tokens are estimated at ~4 chars/token from prompt + response; powers GET /api/usage
# and the "AI TOKENS" widgets in the Sidebar / account dropdown.
USAGE_FILE = os.path.join(os.path.dirname(__file__), "usage.json")
USAGE_QUOTA = 500_000

def _read_usage():
    try:
        with open(USAGE_FILE) as f:
            return json.load(f)
    except Exception:
        return {"used_tokens": 0}

def _record_usage(prompt, result):
    data = _read_usage()
    data["used_tokens"] = data.get("used_tokens", 0) + int((len(prompt) + len(str(result))) / 4)
    try:
        with open(USAGE_FILE, "w") as f:
            json.dump(data, f)
    except Exception:
        pass  # accounting must never break a generation


# Run a one-shot Gemini generation with a HARD timeout. On timeout/error the caller falls
# back to the offline NLP engine, so the app never hangs on a slow/rate-limited AI call.
async def agent_run(config, prompt, structured=True, timeout=30):
    # Local-LLM engine: the UI sets the key to "LOCAL", so every generator that builds a
    # LocalAgentConfig(api_key="LOCAL", ...) is transparently routed to the local model here.
    # On any failure the caller's except-branch falls back to the offline NLP engine.
    if getattr(config, "api_key", None) and config.api_key.startswith("LOCAL"):
        model_name = config.api_key.split(":", 1)[1] if ":" in config.api_key else None
        from nlp.local_llm import local_generate
        schema = getattr(config, "response_schema", None)
        result = await local_generate(prompt, structured=structured, schema=schema,
                                      model_name=model_name,
                                      timeout=max(timeout, 120))
        _record_usage(prompt, result)
        return result

    async def _call():
        async with Agent(config) as agent:
            response = await agent.chat(prompt)
            return await (response.structured_output() if structured else response.text())
    result = await asyncio.wait_for(_call(), timeout=timeout)
    _record_usage(prompt, result)
    return result

# Load env vars from parent directory .env if it exists
load_dotenv(os.path.join(os.path.dirname(__file__), '..', '.env'))

# Raw uploaded PDFs are kept here so the UI can render the real document (a true PDF
# viewer) instead of only its extracted text. Served by GET /api/documents/{id}/file.
from fastapi.responses import FileResponse
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")

import sys
from pathlib import Path
from sqlalchemy.orm import Session
from db import models, crud
from db.database import engine, get_db

# Startup schema-patching routine
def patch_database_schema(engine):
    from sqlalchemy import inspect, text
    inspector = inspect(engine)
    
    with engine.begin() as conn:
        if "users" in inspector.get_table_names():
            cols = [c["name"] for c in inspector.get_columns("users")]
            if "status" not in cols:
                conn.execute(text("ALTER TABLE users ADD COLUMN status VARCHAR DEFAULT 'free'"))

        if "chat_messages" in inspector.get_table_names():
            cols = [c["name"] for c in inspector.get_columns("chat_messages")]
            if "document_id" not in cols:
                conn.execute(text("ALTER TABLE chat_messages ADD COLUMN document_id INTEGER REFERENCES documents(id)"))
                
        if "artifacts" in inspector.get_table_names():
            cols = [c["name"] for c in inspector.get_columns("artifacts")]
            if "document_id" not in cols:
                conn.execute(text("ALTER TABLE artifacts ADD COLUMN document_id INTEGER REFERENCES documents(id)"))
            if "owner_email" not in cols:
                conn.execute(text("ALTER TABLE artifacts ADD COLUMN owner_email VARCHAR"))
                
        if "flashcards" in inspector.get_table_names():
            cols = [c["name"] for c in inspector.get_columns("flashcards")]
            if "project_id" not in cols:
                conn.execute(text("ALTER TABLE flashcards ADD COLUMN project_id INTEGER REFERENCES projects(id)"))
            if "document_id" not in cols:
                conn.execute(text("ALTER TABLE flashcards ADD COLUMN document_id INTEGER REFERENCES documents(id)"))

        if "project_members" in inspector.get_table_names():
            cols = [c["name"] for c in inspector.get_columns("project_members")]
            if "document_id" not in cols:
                conn.execute(text("ALTER TABLE project_members ADD COLUMN document_id INTEGER REFERENCES documents(id)"))

        if "project_invites" in inspector.get_table_names():
            cols = [c["name"] for c in inspector.get_columns("project_invites")]
            if "document_id" not in cols:
                conn.execute(text("ALTER TABLE project_invites ADD COLUMN document_id INTEGER REFERENCES documents(id)"))

        if "roadmap_items" in inspector.get_table_names():
            cols = [c["name"] for c in inspector.get_columns("roadmap_items")]
            if "active" not in cols:
                conn.execute(text("ALTER TABLE roadmap_items ADD COLUMN active INTEGER DEFAULT 0"))

patch_database_schema(engine)
# Create DB tables
models.Base.metadata.create_all(bind=engine)

# --- PyInstaller Runtime Bootstrap ---
def setup_runtime_paths():
    if getattr(sys, "frozen", False):
        base = Path(sys._MEIPASS)
        # NLTK bundle path
        nltk_dir = base / "nltk_data"
        os.environ["NLTK_DATA"] = str(nltk_dir)
        try:
            import nltk
            nltk.data.path.insert(0, str(nltk_dir))
        except ImportError:
            pass
        return base
    return Path(__file__).resolve().parent.parent

BASE_DIR = setup_runtime_paths()

app = FastAPI()

# Allow the local frontend (Vite dev server on :5173 / Electron) to call the
# backend cross-origin at http://127.0.0.1:8000. Wildcard origin with
# allow_credentials=False (browsers reject "*" + credentials).
from fastapi.middleware.cors import CORSMiddleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Check for API key
api_key = os.getenv("GEMINI_API_KEY") or os.getenv("VITE_GEMINI_API_KEY")

def get_api_key(x_api_key: str | None) -> str | None:
    # Engine toggle from the UI (Settings → AI Engine):
    #   "OFFLINE" -> force the deterministic local NLP/algorithm engine (returns None).
    #   "LOCAL"   -> force the local downloaded LLM (Ollama); routed inside agent_run.
    if x_api_key == "OFFLINE":
        return None
    if x_api_key and x_api_key.startswith("LOCAL"):
        return x_api_key
    # Accept both Google AI Studio keys (AIzaSy...) and Antigravity keys (AQ...).
    if x_api_key and x_api_key.startswith(("AIzaSy", "AQ")):
        return x_api_key
    return api_key

# --- Schemas ---
class ChatRequest(BaseModel):
    message: str | None = None
    content: str | None = None
    role: str | None = None
    context: str | None = None
    api_key: str | None = None
    conversation_id: str = "default-session"
    project_id: int | None = None
    document_id: int | None = None
    persona: str | None = "friendly"

class UserLogin(BaseModel):
    email: str
    name: str
    picture: str | None = None

class FolderCreate(BaseModel):
    name: str

class ProjectCreate(BaseModel):
    name: str
    description: str | None = None
    folder_id: int | None = None

class ProjectInviteReq(BaseModel):
    email: str
    role: str

class UrlIngestRequest(BaseModel):
    url: str
    project_id: int | None = None

class FlashcardRequest(BaseModel):
    topic_or_text: str
    project_id: int | None = None
    document_id: int | None = None
    api_key: str | None = None

class Flashcard(BaseModel):
    id: int | None = None
    front: str
    back: str
    interval: int | None = 1
    ease: float | None = 2.5
    repetitions: int | None = 0
    due_date: str | None = None

class FlashcardList(BaseModel):
    flashcards: list[Flashcard]

class LearningPathRequest(BaseModel):
    topic: str

class ModuleTopic(BaseModel):
    title: str
    description: str
    estimated_time: str

class LearningPathModule(BaseModel):
    title: str
    topics: list[ModuleTopic]

class LearningPathSchema(BaseModel):
    title: str
    description: str
    modules: list[LearningPathModule]

class SearchMaterialItem(BaseModel):
    title: str
    url: str
    snippet: str
    type: str
    relevancy_score: int

class SearchMaterialSchema(BaseModel):
    results: list[SearchMaterialItem]

class SearchRequest(BaseModel):
    query: str

class QuizOption(BaseModel):
    id: str
    text: str

class QuizQuestion(BaseModel):
    question: str
    options: list[QuizOption]
    correct_option_id: str
    answer: str | None = None
    explanation: str

class QuizSchema(BaseModel):
    title: str
    questions: list[QuizQuestion]

class NoteSection(BaseModel):
    heading: str
    bullet_points: list[str]

class NoteSchema(BaseModel):
    title: str
    summary: str
    sections: list[NoteSection]

class TopicRequest(BaseModel):
    topic_or_text: str
    api_key: str | None = None
    project_id: int | None = None
    document_id: int | None = None
    page_ranges: list[int] | None = None
    
class ReviewRequest(BaseModel):
    id: int | None = None
    interval: int
    ease: float
    repetitions: int
    due: str
    quality: int


class SuggestionSchema(BaseModel):
    path_topic: str
    quiz_topic: str
    flashcard_topic: str

class ConceptNodeSchema(BaseModel):
    id: str
    label: str
    definition: str = Field(description="A clear and concise definition or explanation of the concept")
    formula: str | None = Field(default=None, description="The mathematical or chemical formula associated with the concept, if applicable")

class ConceptEdgeSchema(BaseModel):
    source: str
    target: str
    weight: int

class ConceptMapSchema(BaseModel):
    nodes: list[ConceptNodeSchema]
    edges: list[ConceptEdgeSchema]

class QuizSubmitRequest(BaseModel):
    document_id: int
    score: int
    total_questions: int

class UpgradeRequest(BaseModel):
    email: str

class RoadmapItemUpdate(BaseModel):
    completed: bool | None = None
    active: bool | None = None

# --- Folder Endpoints ---
@app.post("/api/user/upgrade")
def upgrade_user(req: UpgradeRequest, db: Session = Depends(get_db)):
    db_user = db.query(models.User).filter(models.User.email == req.email).first()
    if not db_user:
        db_user = models.User(email=req.email, name=req.email.split('@')[0], status="premium")
        db.add(db_user)
    else:
        db_user.status = "premium"
    db.commit()
    db.refresh(db_user)
    return {"status": "premium", "email": req.email}

@app.post("/api/auth/login")
def login(user_data: UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(models.User).filter(models.User.email == user_data.email).first()
    if db_user:
        db_user.last_login = datetime.datetime.utcnow()
        db_user.name = user_data.name
        if user_data.picture:
            db_user.picture = user_data.picture
        db.commit()
        db.refresh(db_user)
        return db_user
    else:
        new_user = models.User(email=user_data.email, name=user_data.name, picture=user_data.picture)
        db.add(new_user)
        db.commit()
        db.refresh(new_user)
        return new_user

@app.get("/api/user/stats")
def get_user_stats(email: str, db: Session = Depends(get_db)):
    db_user = db.query(models.User).filter(models.User.email == email).first()
    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")
        
    doc_count = db.query(models.Document).count() # In a real app this would filter by user_id
    project_count = db.query(models.Project).count() # Same here
    
    return {
        "document_count": doc_count,
        "project_count": project_count,
        "last_login": db_user.last_login.isoformat() if db_user.last_login else None
    }

@app.get("/api/folders")
async def get_folders_api(db: Session = Depends(get_db)):
    folders = crud.get_folders(db)
    return [{"id": f.id, "name": f.name} for f in folders]

@app.post("/api/folders")
async def create_folder_api(folder: FolderCreate, db: Session = Depends(get_db)):
    return crud.create_folder(db, name=folder.name)

# --- Project Endpoints ---

@app.post("/api/projects")
async def create_project_api(project: ProjectCreate, db: Session = Depends(get_db)):
    return crud.create_project(db, name=project.name, description=project.description, folder_id=project.folder_id)

@app.get("/api/projects")
async def get_projects_api(folder_id: int | None = None, db: Session = Depends(get_db)):
    projects = crud.get_projects(db, folder_id=folder_id)
    result = []
    for p in projects:
        result.append({
            "id": p.id,
            "name": p.name,
            "description": p.description,
            "folder_id": p.folder_id,
            "created_at": p.created_at,
            "source_count": len(p.documents)
        })
    return result

@app.get("/api/projects/{project_id}")
async def get_project_api(project_id: int, db: Session = Depends(get_db)):
    project = crud.get_project(db, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return {
        "id": project.id,
        "name": project.name,
        "description": project.description,
        "created_at": project.created_at,
        "documents": [{"id": d.id, "filename": d.filename, "kind": d.kind, "upload_date": d.upload_date} for d in project.documents],
        "messages": [{"id": m.id, "role": m.role, "content": m.content, "persona": m.persona, "created_at": m.created_at} for m in project.messages]
    }

@app.post("/api/projects/{project_id}/invite")
def invite_member(project_id: int, req: ProjectInviteReq, db: Session = Depends(get_db)):
    proj = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    
    existing_invite = db.query(models.ProjectInvite).filter(
        models.ProjectInvite.project_id == project_id,
        models.ProjectInvite.email == req.email
    ).first()
    
    if existing_invite:
        return {"status": "Already invited"}
        
    new_invite = models.ProjectInvite(project_id=project_id, email=req.email, role=req.role)
    db.add(new_invite)
    db.commit()
    return {"status": "Invited successfully"}

@app.get("/api/projects/{project_id}/members")
def get_project_members(project_id: int, db: Session = Depends(get_db)):
    project = crud.get_project(db, project_id=project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    # In a local app, the owner is the current user. Since we don't have sessions, we mock the owner
    members = db.query(models.ProjectMember).filter(models.ProjectMember.project_id == project_id).all()
    invites = db.query(models.ProjectInvite).filter(models.ProjectInvite.project_id == project_id).all()
    
    # Return mock owner if members is empty (legacy projects)
    if not members:
        members_data = [{"email": "owner@local.app", "role": "owner"}]
    else:
        members_data = [{"email": m.email, "role": m.role} for m in members]
        
    return {
        "active": members_data,
        "pending": [{"email": i.email, "role": i.role, "status": i.status} for i in invites]
    }

@app.delete("/api/projects/{project_id}")
async def delete_project_api(project_id: int, db: Session = Depends(get_db)):
    success = crud.delete_project(db, project_id=project_id)
    if not success:
        raise HTTPException(status_code=404, detail="Project not found")
    return {"status": "ok"}

@app.get("/api/projects/{project_id}/roadmap")
async def get_project_roadmap(project_id: int, db: Session = Depends(get_db)):
    roadmap = crud.get_roadmap(db, project_id=project_id)
    if not roadmap:
        return {"items": []}
    return {"items": [{"id": i.id, "step_number": i.step_number, "title": i.title, "description": i.description, "completed": i.completed, "active": i.active} for i in sorted(roadmap.items, key=lambda x: x.step_number)]}

@app.post("/api/projects/{project_id}/roadmap/generate")
async def generate_project_roadmap(project_id: int, req: TopicRequest, db: Session = Depends(get_db)):
    project = crud.get_project(db, project_id=project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    # 1. Gather all document texts
    docs = crud.get_documents(db, project_id=project_id)
    text_content = ""
    for doc in docs:
        if doc.content:
            text_content += doc.content + "\n\n"
    if not text_content.strip():
        text_content = "Vui lòng cung cấp nội dung chung."

    # 2. Call roadmap generator
    from nlp.roadmap import generate_roadmap
    items = await generate_roadmap(text_content, api_key=req.api_key)
    
    # 3. Save to DB
    crud.create_roadmap(db, project_id=project_id, items=items)
    return {"status": "ok", "message": "Đã tạo giáo án"}

@app.get("/api/documents/{document_id}/roadmap")
async def get_document_roadmap(document_id: int, db: Session = Depends(get_db)):
    roadmap = db.query(models.Roadmap).filter(models.Roadmap.document_id == document_id).first()
    if not roadmap:
        return {"items": []}
    return {"items": [{"id": i.id, "step_number": i.step_number, "title": i.title, "description": i.description, "completed": i.completed, "active": i.active} for i in sorted(roadmap.items, key=lambda x: x.step_number)]}

@app.patch("/api/roadmap/items/{item_id}")
async def update_roadmap_item_endpoint(item_id: int, req: RoadmapItemUpdate, db: Session = Depends(get_db)):
    item = crud.update_roadmap_item(db, item_id=item_id, completed=req.completed, active=req.active)
    if not item:
        raise HTTPException(status_code=404, detail="Roadmap item not found")
    return {
        "id": item.id,
        "step_number": item.step_number,
        "title": item.title,
        "description": item.description,
        "completed": item.completed,
        "active": item.active
    }

@app.post("/api/documents/{document_id}/roadmap/generate")
async def generate_document_roadmap(document_id: int, req: TopicRequest, db: Session = Depends(get_db)):
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc or not doc.content:
        return {"status": "error", "message": "Không tìm thấy nội dung tài liệu."}
    from nlp.roadmap import generate_roadmap
    items = await generate_roadmap(doc.content, api_key=req.api_key)
    
    old_roadmaps = db.query(models.Roadmap).filter(models.Roadmap.document_id == document_id).all()
    for rm in old_roadmaps:
        db.query(models.RoadmapItem).filter(models.RoadmapItem.roadmap_id == rm.id).delete()
        db.delete(rm)
    db.commit()
    
    db_roadmap = models.Roadmap(document_id=document_id)
    db.add(db_roadmap)
    db.commit()
    db.refresh(db_roadmap)
    for i, item in enumerate(items):
        db_item = models.RoadmapItem(
            roadmap_id=db_roadmap.id,
            step_number=i + 1,
            title=item.get("title", ""),
            description=item.get("description", "")
        )
        db.add(db_item)
    db.commit()
    return {"status": "ok", "message": "Đã tạo giáo án cho tài liệu"}

@app.get("/api/documents/{document_id}/members")
def get_document_members(document_id: int, db: Session = Depends(get_db)):
    members = db.query(models.ProjectMember).filter(models.ProjectMember.document_id == document_id).all()
    invites = db.query(models.ProjectInvite).filter(models.ProjectInvite.document_id == document_id).all()
    if not members:
        members_data = [{"email": "owner@local.app", "role": "owner"}]
    else:
        members_data = [{"email": m.email, "role": m.role} for m in members]
    return {
        "active": members_data,
        "pending": [{"email": i.email, "role": i.role, "status": i.status} for i in invites]
    }

@app.post("/api/documents/{document_id}/invite")
def invite_document_member(document_id: int, req: ProjectInviteReq, db: Session = Depends(get_db)):
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    exist = db.query(models.ProjectInvite).filter(models.ProjectInvite.document_id == document_id, models.ProjectInvite.email == req.email).first()
    if exist:
        return {"status": "Already invited"}
    invite = models.ProjectInvite(document_id=document_id, email=req.email, role=req.role)
    db.add(invite)
    db.commit()
    return {"status": "Invited"}

@app.get("/api/projects/{project_id}/artifacts")
async def get_project_artifacts(project_id: int, db: Session = Depends(get_db)):
    artifacts = crud.get_artifacts(db, project_id=project_id)
    return [{"id": a.id, "type": a.type, "title": a.title, "content": a.content, "created_at": a.created_at} for a in artifacts]

@app.get("/api/documents/{document_id}/artifacts")
async def get_document_artifacts(document_id: int, db: Session = Depends(get_db)):
    artifacts = crud.get_artifacts(db, document_id=document_id)
    return [{"id": a.id, "type": a.type, "title": a.title, "content": a.content, "created_at": a.created_at} for a in artifacts]

@app.get("/api/documents/{document_id}/notes")
async def get_document_notes_api(document_id: int, db: Session = Depends(get_db)):
    # Check if artifact exists
    artifact = db.query(models.Artifact).filter(models.Artifact.document_id == document_id, models.Artifact.type == 'smart_notes').first()
    if artifact:
        import json
        return json.loads(artifact.content)
    
    # Generate notes if not exists
    text = get_document_text(db, document_id)
    from nlp.notes import extract_notes
    notes_dict = extract_notes(text)
    
    # Save it
    import json
    new_artifact = models.Artifact(
        document_id=document_id,
        type='smart_notes',
        title=notes_dict.get('title', 'Smart Notes'),
        content=json.dumps(notes_dict, ensure_ascii=False)
    )
    db.add(new_artifact)
    db.commit()
    
    return notes_dict

# --- AI Engine ---

@app.get("/api/engine/local_status")
async def engine_local_status(model_name: str | None = None):
    """Whether a local downloaded model (Ollama) is reachable, for the Settings checker."""
    from nlp.local_llm import status
    return status(model_name)

@app.get("/api/engine/recommend_local")
async def recommend_local_model():
    """Detect system RAM and recommend the best local model for Vietnamese & study app."""
    import os
    ram_gb = 8.0 # default
    try:
        if hasattr(os, 'sysconf'):
            if 'SC_PAGE_SIZE' in os.sysconf_names and 'SC_PHYS_PAGES' in os.sysconf_names:
                ram_gb = (os.sysconf('SC_PAGE_SIZE') * os.sysconf('SC_PHYS_PAGES')) / (1024.**3)
    except Exception:
        pass
    
    ram_gb = round(ram_gb, 1)
    
    if ram_gb >= 14:
        rec = "gemma2"
        reason = f"Máy bạn có {ram_gb}GB RAM, cực kỳ mạnh mẽ để chạy Gemma 2 (9B) hoặc Qwen2.5 (14B) - các model vượt trội về khả năng suy luận logic và tiếng Việt."
        alts = ["qwen2.5:14b", "llama3.1", "qwen2.5"]
    elif ram_gb >= 6:
        rec = "qwen2.5"
        reason = f"Máy bạn có {ram_gb}GB RAM, đủ để chạy mượt Qwen2.5 (7B) hoặc Llama 3.1 (8B) - cân bằng hoàn hảo giữa tốc độ và chất lượng tiếng Việt."
        alts = ["llama3.1", "gemma2", "mistral"]
    else:
        rec = "qwen2.5:3b"
        reason = f"Máy bạn có {ram_gb}GB RAM, phù hợp chạy các model nhỏ và nhẹ như Qwen2.5 (3B) hoặc Gemma 2 (2B) để đảm bảo không bị giật lag."
        alts = ["gemma2:2b", "llama3.2"]

    return {
        "system_ram_gb": ram_gb,
        "recommended_model": rec,
        "reasoning": reason,
        "alternatives": alts
    }

@app.get("/api/usage")
async def get_usage():
    """AI-token usage for the Sidebar / account-dropdown widgets (estimated, local-only)."""
    used = _read_usage().get("used_tokens", 0)
    remaining = max(USAGE_QUOTA - used, 0)
    percent = round(used / USAGE_QUOTA * 100, 2) if USAGE_QUOTA else 0.0
    return {"used": used, "quota": USAGE_QUOTA, "remaining": remaining, "percent": percent}

# --- Endpoints ---

def offline_doc_answer(message: str, db: Session, context: str = "") -> str:
    """Offline chat: answer from the active (context) or latest document via TF-IDF sentence retrieval."""
    try:
        source_text = context.strip() if context and len(context.strip()) > 30 else ""
        label = "tài liệu đang chọn"
        if not source_text:
            doc = db.query(models.Document).order_by(models.Document.id.desc()).first()
            if not doc or not doc.content or not doc.content.strip():
                return "📂 [Offline] Chưa có tài liệu nào trong thư viện. Hãy Upload một tài liệu để mình trả lời dựa trên nội dung của nó."
            source_text = doc.content
            label = doc.filename
            
        import re
        from nlp.vietnamese import VIETNAMESE_STOPWORDS, is_vietnamese
        
        # Split by punctuation OR by newlines to handle bullet points and math cheatsheets better
        sentences = [s.strip() for s in re.split(r'(?<=[.!?…])\s+|\n+', source_text) if len(s.strip()) > 3]
        if not sentences:
            return "📄 [Offline] Tài liệu hiện chưa có nội dung văn bản để trích xuất."
            
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        
        stop_words = list(VIETNAMESE_STOPWORDS) if is_vietnamese(source_text) else "english"
        
        # For answering questions
        if message and message.strip() and not any(kw in message.lower() for kw in ["tóm tắt", "summary", "summarize"]):
            try:
                matrix = TfidfVectorizer(stop_words=stop_words).fit_transform(sentences + [message])
                sims = cosine_similarity(matrix[-1], matrix[:-1]).flatten()
                ranked = [int(i) for i in sims.argsort()[::-1] if sims[int(i)] > 0.01][:4]
                if ranked:
                    answer = "\n".join("- " + sentences[i] for i in ranked)
                    return f"📚 [Offline · trích từ {label}]\n\n{answer}"
            except Exception:
                pass
                
        # Fall back to an extractive summary using TF-IDF sentence centrality.
        try:
            doc_matrix = TfidfVectorizer(stop_words=stop_words).fit_transform(sentences)
            centrality = (doc_matrix * doc_matrix.T).toarray().sum(axis=1)
            top = sorted(sorted(range(len(sentences)), key=lambda i: centrality[i], reverse=True)[:5])
            summary = "\n".join("- " + sentences[i] for i in top)
            return f"📝 [Offline · tóm tắt {label}]\n\n{summary}"
        except Exception:
            # If TF-IDF fails (e.g. not enough words), just return the first few sentences
            return f"📝 [Offline · tóm tắt {label}]\n\n" + "\n".join("- " + s for s in sentences[:5])
            
    except Exception as e:
        return f"⚠️ [Offline] Lỗi khi truy xuất tài liệu: {e}"


PERSONA_PROMPTS = {
    "Gia sư thân thiện": "You are Workflow Assistant, a friendly tutor. Giảng bài cặn kẽ, rõ ràng, luôn khích lệ, tóm tắt các điểm chính. Trả lời bằng tiếng Việt hoặc tiếng Anh tuỳ ngôn ngữ của học sinh.",
    "Coach nghiêm túc": "You are Workflow Coach, a strict and direct mentor. Dùng từ ngữ ngắn gọn, đi thẳng vào vấn đề, yêu cầu học sinh làm bài tập. Không giải thích dông dài. Trả lời bằng tiếng Việt hoặc tiếng Anh tuỳ ngôn ngữ của học sinh.",
    "Socratic hỏi gợi mở": "You are Socratic Workflow, a mentor. Dẫn dắt bằng câu hỏi gợi mở, không bật đáp án ngay. Luôn hỏi 1-2 câu cho học sinh tự suy luận trước. Trả lời bằng tiếng Việt hoặc tiếng Anh tuỳ ngôn ngữ của học sinh.",
    "Bạn học Gen Z": "You are Workflow Buddy, a Gen-Z peer. Trẻ trung, năng lượng, dùng bullet ngắn gọn, ví dụ vui vẻ dễ nhớ, luôn có bước tiếp theo. Trả lời bằng tiếng Việt hoặc tiếng Anh tuỳ ngôn ngữ của học sinh."
}

@app.post("/api/chat")
async def chat_endpoint(request: ChatRequest, db: Session = Depends(get_db)):
    message_text = request.message or request.content
    if not message_text:
        raise HTTPException(status_code=400, detail="Message/content required")
        
    role = request.role or "user"
    
    if request.project_id or request.document_id:
        crud.save_chat_message(db, project_id=request.project_id, role=role, content=message_text, persona=request.persona, document_id=request.document_id)

    context = request.context or ""
    if request.project_id and not context:
        project = crud.get_project(db, request.project_id)
        if project:
            sources = [doc.content for doc in project.documents if doc.content]
            context = "\n\n---\n\n".join(sources)[:30000]
    elif request.document_id and not context:
        doc = db.query(models.Document).filter(models.Document.id == request.document_id).first()
        if doc and doc.content:
            context = doc.content[:30000]
            
    # persona=None means "Chat chính" (plain AI Q&A, no tutor persona).
    if request.persona is None:
        persona_prompt = "Bạn là trợ lý AI trả lời trực tiếp, chính xác và ngắn gọn dựa trên tài liệu."
    else:
        persona_prompt = PERSONA_PROMPTS.get(request.persona, PERSONA_PROMPTS["Gia sư thân thiện"])
    
    active_step_context = ""
    if request.project_id:
        rm = db.query(models.Roadmap).filter(models.Roadmap.project_id == request.project_id).first()
        if rm:
            active_item = db.query(models.RoadmapItem).filter(models.RoadmapItem.roadmap_id == rm.id, models.RoadmapItem.active == 1).first()
            if active_item:
                active_step_context = f"Chủ đề người học đang tập trung trong lộ trình: '{active_item.title}' - Mô tả: '{active_item.description}'."
    elif request.document_id:
        rm = db.query(models.Roadmap).filter(models.Roadmap.document_id == request.document_id).first()
        if rm:
            active_item = db.query(models.RoadmapItem).filter(models.RoadmapItem.roadmap_id == rm.id, models.RoadmapItem.active == 1).first()
            if active_item:
                active_step_context = f"Chủ đề người học đang tập trung trong lộ trình: '{active_item.title}' - Mô tả: '{active_item.description}'."

    if active_step_context:
        persona_prompt = f"{active_step_context}\n{persona_prompt}"
    
    current_key = get_api_key(request.api_key)
    response_text = ""
    if current_key:
        try:
            config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite", capabilities=FAST_CAPS)
            prompt = (
                f"{persona_prompt}\n\nUsing the document context below, answer the question.\n"
                f"Context:\n{context}\n\nQuestion: {message_text}"
            )
            response_text = await agent_run(config, prompt, structured=False)
        except Exception as e:
            print(f"[chat] AI failed, falling back to offline: {e}")
            
    if not response_text:
        response_text = offline_doc_answer(message_text, db, context)

    msg_id = None
    if request.project_id or request.document_id:
        msg = crud.save_chat_message(db, project_id=request.project_id, role="assistant", content=response_text, persona=request.persona, document_id=request.document_id)
        msg_id = msg.id

    return {"id": msg_id, "response": response_text}

@app.get("/api/projects/{project_id}/messages")
async def get_project_messages_api(project_id: int, db: Session = Depends(get_db)):
    messages = crud.get_chat_history(db, project_id=project_id)
    return {
        "messages": [
            {
                "id": m.id, "role": m.role, "content": m.content, 
                "persona": m.persona, "created_at": m.created_at.isoformat() if m.created_at else None
            } for m in messages
        ]
    }

@app.get("/api/documents/{document_id}/messages")
async def get_document_messages_api(document_id: int, db: Session = Depends(get_db)):
    messages = crud.get_chat_history(db, document_id=document_id)
    return {
        "messages": [
            {
                "id": m.id, "role": m.role, "content": m.content, 
                "persona": m.persona, "created_at": m.created_at.isoformat() if m.created_at else None
            } for m in messages
        ]
    }

@app.post("/api/generate_flashcards")
async def generate_flashcards(request: FlashcardRequest, db: Session = Depends(get_db)):
    text_to_process = request.topic_or_text
    if request.document_id:
        doc_text = get_document_text(db, request.document_id)
        if doc_text.strip():
            text_to_process = doc_text

    def _save(cards):
        if request.project_id or request.document_id:
            query = db.query(models.Flashcard)
            if request.document_id:
                query = query.filter(models.Flashcard.document_id == request.document_id)
            elif request.project_id:
                query = query.filter(models.Flashcard.project_id == request.project_id)
            query.delete()
            db.commit()

        out = []
        for c in cards:
            db_card = crud.create_flashcard(db, front=c.get("front", ""), back=c.get("back", ""), project_id=request.project_id, document_id=request.document_id)
            out.append({
                "id": db_card.id, "front": db_card.front, "back": db_card.back,
                "interval": db_card.interval, "ease": db_card.ease,
                "repetitions": db_card.repetitions, "due_date": db_card.due_date.isoformat()
            })
        return out

    current_key = get_api_key(request.api_key)
    if current_key:
        try:
            config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite", response_schema=FlashcardList, capabilities=FAST_CAPS)
            prompt = f"Generate 5 to 10 highly effective study flashcards based on the following topic or text. Front = a clear question or concept, back = a concise answer or definition. Use the SAME language as the text.\n\nText:\n{text_to_process}"
            data = await agent_run(config, prompt)
            if data and data.get("flashcards"):
                return {"flashcards": _save(data["flashcards"])}
        except Exception as e:
            print(f"[flashcards] AI failed, falling back to offline: {e}")
    # Offline NLP engine (default + fallback)
    from nlp.flashcards import extract_flashcards
    return {"flashcards": _save(extract_flashcards(text_to_process, max_cards=10))}

from fastapi import Form

@app.post("/api/documents/upload")
def upload_document(
    file: UploadFile = File(...),
    project_id: int | None = Form(None),
    db: Session = Depends(get_db)
):
    import tempfile, os, shutil
    
    suffix = os.path.splitext(file.filename)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        path = tmp.name
    
    try:
        text = ""
        kind = "document"
        page_count = None
        pages_data = None
        
        if suffix == ".pdf":
            kind = "pdf"
            import pdfplumber
            try:
                with pdfplumber.open(path) as pdf:
                    pages = []
                    for page in pdf.pages:
                        page_text = (page.extract_text() or "") + "\n"
                        pages.append(page_text)
                        text += page_text
                    page_count = len(pages)
                    pages_data = json.dumps(pages)
            except Exception:
                pass
        elif suffix in [".docx"]:
            kind = "docx"
            try:
                import docx
                d = docx.Document(path)
                text = "\n".join([p.text for p in d.paragraphs])
            except Exception:
                pass
        elif suffix in [".png", ".jpg", ".jpeg", ".webp"]:
            kind = "image"
            try:
                import pytesseract
                from PIL import Image
                text = pytesseract.image_to_string(Image.open(path), lang="vie+eng")
            except Exception:
                pass
        elif suffix in [".mp4", ".mp3", ".m4a", ".wav"]:
            kind = "audio" if suffix in [".mp3", ".m4a", ".wav"] else "video"
            try:
                from faster_whisper import WhisperModel
                model = WhisperModel("tiny", device="cpu", compute_type="int8")
                segments, info = model.transcribe(path)
                text = "\n".join([segment.text for segment in segments])
            except Exception:
                pass
        else:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read()
            except Exception:
                pass
                
        doc = crud.create_document(db, filename=file.filename, content=text, kind=kind, project_id=project_id, page_count=page_count, pages_data=pages_data)
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        shutil.copyfile(path, os.path.join(UPLOAD_DIR, f"{doc.id}{suffix}"))
        return {"id": doc.id, "filename": doc.filename, "kind": doc.kind, "upload_date": doc.upload_date.isoformat()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        os.unlink(path)

@app.post("/api/documents/ingest_url")
@app.post("/api/documents/url")
def ingest_url(request: UrlIngestRequest, db: Session = Depends(get_db)):
    url = request.url
    text = ""
    filename = url
    kind = "link"
    
    if "youtube.com" in url or "youtu.be" in url:
        kind = "video"
        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            import urllib.parse
            if "youtu.be" in url:
                video_id = url.split("/")[-1].split("?")[0]
            else:
                parsed = urllib.parse.urlparse(url)
                video_id = urllib.parse.parse_qs(parsed.query).get('v', [None])[0]
            if video_id:
                try:
                    transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['vi', 'en'])
                    text = "\n".join([t['text'] for t in transcript])
                    filename = f"YouTube: {video_id}"
                except Exception:
                    pass
        except ImportError:
            pass
    else:
        try:
            import trafilatura
            downloaded = trafilatura.fetch_url(url)
            if downloaded:
                text = trafilatura.extract(downloaded) or ""
        except ImportError:
            pass
            
    doc = crud.create_document(db, filename=filename, content=text, kind=kind, project_id=request.project_id)
    return {"id": doc.id, "filename": doc.filename, "kind": doc.kind, "upload_date": doc.upload_date.isoformat()}

@app.get("/api/documents")
async def list_documents(project_id: int = None, db: Session = Depends(get_db)):
    docs = crud.get_documents(db, project_id=project_id)
    import glob
    def has_file(doc):
        suffix = os.path.splitext(doc.filename)[1].lower()
        file_path = os.path.join(UPLOAD_DIR, f"{doc.id}{suffix}")
        if os.path.exists(file_path):
            return True
        files = [f for f in glob.glob(os.path.join(UPLOAD_DIR, f"{doc.id}.*"))]
        if suffix != '.txt':
            files = [f for f in files if not f.endswith('.txt')]
        return len(files) > 0
    return {"documents": [{"id": d.id, "filename": d.filename, "kind": d.kind, "upload_date": d.upload_date.isoformat(), "content": d.content,
                           "project_id": d.project_id, "has_file": has_file(d)} for d in docs]}

@app.get("/api/documents/{document_id}")
async def get_document(document_id: int, db: Session = Depends(get_db)):
    """Single document with per-page text — powers the exam 'NGUỒN: TÀI LIỆU ĐỐI CHIẾU'
    reference panel (page-by-page navigation)."""
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    pages = []
    if doc.pages_data:
        try:
            pages = json.loads(doc.pages_data)
        except Exception:
            pages = []
    if not pages and doc.content:
        pages = [doc.content]
    return {
        "id": doc.id,
        "filename": doc.filename,
        "kind": doc.kind,
        "page_count": doc.page_count or len(pages),
        "content": doc.content,
        "pages": pages,
    }

@app.get("/api/documents/{document_id}/file")
async def get_document_file(document_id: int, db: Session = Depends(get_db)):
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    suffix = os.path.splitext(doc.filename)[1].lower()
    file_path = os.path.join(UPLOAD_DIR, f"{doc.id}{suffix}")
    if not os.path.exists(file_path):
        import glob
        files = [f for f in glob.glob(os.path.join(UPLOAD_DIR, f"{document_id}.*"))]
        if suffix != '.txt':
            files = [f for f in files if not f.endswith('.txt')]
        if not files:
            raise HTTPException(status_code=404, detail="No source file stored for this document")
        file_path = files[0]
        
    ext = os.path.splitext(file_path)[1].lower()
    media_types = {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".mp4": "video/mp4",
        ".mp3": "audio/mpeg",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain"
    }
    media_type = media_types.get(ext, "application/octet-stream")
    return FileResponse(
        file_path, 
        media_type=media_type, 
        filename=os.path.basename(file_path),
        content_disposition_type="inline"
    )

@app.delete("/api/documents/{document_id}")
async def remove_document(document_id: int, db: Session = Depends(get_db)):
    ok = crud.delete_document(db, document_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"success": True, "deleted_id": document_id}

@app.post("/api/generate_path")
async def generate_path(request: LearningPathRequest, x_api_key: str | None = Header(default=None), db: Session = Depends(get_db)):
    current_key = get_api_key(x_api_key)
    if current_key:
        try:
            config = LocalAgentConfig(
                api_key=current_key,
                response_schema=LearningPathSchema
            )
            prompt = f"Create a comprehensive, step-by-step learning path for the topic: '{request.topic}'. Break it down into logical modules, and each module should have specific topics with descriptions and estimated study times."
            
            async with Agent(config) as agent:
                response = await agent.chat(prompt)
                data = await response.structured_output()
                if data:
                    return data
        except Exception as e:
            print(f"Error generating learning path, falling back to offline: {e}")

    from nlp.concept_map import generate_offline_learning_path
    return generate_offline_learning_path(request.topic, db=db)

import json
from ddgs import DDGS

@app.post("/api/search")
async def search_materials(request: SearchRequest, x_api_key: str | None = Header(default=None)):
    if not request.query or not request.query.strip():
        return {"results": []}
    current_key = get_api_key(x_api_key)
    try:
        raw_results = []
        with DDGS() as ddgs:
            ddgs_results = list(ddgs.text(f"{request.query} guide OR tutorial OR pdf OR course", max_results=15))
            for res in ddgs_results:
                raw_results.append({
                    "title": res.get('title', ''),
                    "url": res.get('href', ''),
                    "snippet": res.get('body', '')
                })
        
        # If we have a valid key, use AI to filter and score for quality
        if current_key:
            config = LocalAgentConfig(
                api_key=current_key,
                response_schema=SearchMaterialSchema
            )
            prompt = f"""
            You are a helpful and intelligent curator. A user is searching for materials on the topic: '{request.query}'.
            I have queried a search engine and retrieved the following raw results:
            {json.dumps(raw_results, indent=2, ensure_ascii=False)}
            
            Your task is to review these raw results, evaluate their relevance and quality, and select the TOP 5 best resources.
            For each selected resource:
            - Keep the EXACT original 'url'. Do not change it.
            - Clean up the 'title' if necessary.
            - Write a brief 'snippet' explaining WHY this is a good resource. Respond in the language of the query.
            - Categorize its 'type' strictly as one of: Video, Article, PDF, Book, Code, or Documentation.
            - Assign a 'relevancy_score' from 0 to 100 indicating how relevant and academically/professionally useful the result is.
            """
            
            async with Agent(config) as agent:
                response = await agent.chat(prompt)
                data = await response.structured_output()
                if data and data.get("results"):
                    return data

        # Fallback to top 5 if AI fails or no key
        fallback_results = []
        for res in raw_results[:5]:
            url = res["url"]
            type_label = "Article"
            if "youtube.com" in url or "youtu.be" in url:
                type_label = "Video"
            elif url.endswith(".pdf"):
                type_label = "PDF"
            elif "amazon.com" in url or "goodreads.com" in url:
                type_label = "Book"
            elif "github.com" in url:
                type_label = "Code"
            
            fallback_results.append({
                "title": res["title"],
                "url": url,
                "snippet": res["snippet"],
                "type": type_label,
                "relevancy_score": 80
            })
            
        return {"results": fallback_results}
            
    except Exception as e:
        print(f"Error searching materials: {e}")
        raise HTTPException(status_code=500, detail=str(e))

def get_document_text(db: Session, document_id: int | None, page_ranges: list[int] | None = None) -> str:
    if not document_id:
        return ""
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        return ""
    if page_ranges:
        try:
            pages = json.loads(doc.pages_data) if doc.pages_data else []
            selected_text = ""
            for p in page_ranges:
                idx = p - 1
                if 0 <= idx < len(pages):
                    selected_text += pages[idx] + "\n"
            if selected_text.strip():
                return selected_text
        except Exception:
            pass
    if doc.content:
        return doc.content
    if doc.pages_data:
        try:
            pages = json.loads(doc.pages_data)
            return "\n".join(pages)
        except Exception:
            pass
    return ""

@app.post("/api/generate_quiz")
async def generate_quiz_endpoint(request: TopicRequest, db: Session = Depends(get_db)):
    text_to_process = request.topic_or_text
    
    if request.document_id:
        doc_text = get_document_text(db, request.document_id, request.page_ranges)
        if doc_text.strip():
            text_to_process = doc_text

    current_key = get_api_key(request.api_key)
    if current_key:
        try:
            config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite", response_schema=QuizSchema, capabilities=FAST_CAPS)
            prompt = f"Generate a multiple-choice quiz (5 questions), in the SAME language as the text: '{text_to_process}'. Each question must have 4 options, a correct option ID, and an explanation."
            data = await agent_run(config, prompt)
            if data and data.get("questions"):
                if request.project_id or request.document_id:
                    query = db.query(models.Artifact).filter(models.Artifact.type == "quiz")
                    if request.document_id:
                        query = query.filter(models.Artifact.document_id == request.document_id)
                    elif request.project_id:
                        query = query.filter(models.Artifact.project_id == request.project_id)
                    query.delete()
                    db.commit()
                    crud.create_artifact(db, project_id=request.project_id, type="quiz", title=data.get("title", "Bài Trắc Nghiệm"), content=json.dumps(data), document_id=request.document_id)
                return data
        except Exception as e:
            print(f"[quiz] AI failed, falling back to offline: {e}")
    from nlp.quizzes import extract_quiz
    quiz_data = extract_quiz(text_to_process, num_questions=5)
    if request.project_id or request.document_id:
        query = db.query(models.Artifact).filter(models.Artifact.type == "quiz")
        if request.document_id:
            query = query.filter(models.Artifact.document_id == request.document_id)
        elif request.project_id:
            query = query.filter(models.Artifact.project_id == request.project_id)
        query.delete()
        db.commit()
        crud.create_artifact(db, project_id=request.project_id, type="quiz", title=quiz_data.get("title", "Bài Trắc Nghiệm"), content=json.dumps(quiz_data), document_id=request.document_id)
    return quiz_data
class ExamPrepSchema(BaseModel):
    title: str
    markdown_content: str

@app.post("/api/generate_exam_prep")
async def generate_exam_prep_endpoint(request: TopicRequest, db: Session = Depends(get_db)):
    text_to_process = request.topic_or_text
    
    if request.document_id:
        doc_text = get_document_text(db, request.document_id, request.page_ranges)
        if doc_text.strip():
            text_to_process = doc_text

    current_key = get_api_key(request.api_key)
    if current_key:
        try:
            config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite", response_schema=ExamPrepSchema, capabilities=FAST_CAPS)
            prompt = f"Generate a comprehensive Exam Preparation cheat sheet or study guide based on the following text: '{text_to_process}'. Use markdown formatting."
            data = await agent_run(config, prompt)
            if data and data.get("markdown_content"):
                if request.project_id or request.document_id:
                    query = db.query(models.Artifact).filter(models.Artifact.type == "examprep")
                    if request.document_id:
                        query = query.filter(models.Artifact.document_id == request.document_id)
                    elif request.project_id:
                        query = query.filter(models.Artifact.project_id == request.project_id)
                    query.delete()
                    db.commit()
                    crud.create_artifact(db, project_id=request.project_id, type="examprep", title=data.get("title", "Tài liệu phòng thi"), content=json.dumps(data), document_id=request.document_id)
                return data
        except Exception as e:
            print(f"[exam_prep] AI failed, falling back to offline: {e}")
            
    # Fallback Offline Engine
    from nlp.concept_map import generate_offline_exam_prep
    data = generate_offline_exam_prep(text_to_process)
    if request.project_id or request.document_id:
        query = db.query(models.Artifact).filter(models.Artifact.type == "examprep")
        if request.document_id:
            query = query.filter(models.Artifact.document_id == request.document_id)
        elif request.project_id:
            query = query.filter(models.Artifact.project_id == request.project_id)
        query.delete()
        db.commit()
        crud.create_artifact(db, project_id=request.project_id, type="examprep", title=data.get("title", "Tài liệu phòng thi"), content=json.dumps(data), document_id=request.document_id)
    return data

class StudyPlanSchema(BaseModel):
    title: str
    markdown_content: str

@app.post("/api/generate_study_plan")
async def generate_study_plan_endpoint(request: TopicRequest, db: Session = Depends(get_db)):
    text_to_process = request.topic_or_text
    if request.document_id:
        doc_text = get_document_text(db, request.document_id, request.page_ranges)
        if doc_text.strip():
            text_to_process = doc_text

    current_key = get_api_key(request.api_key)
    if current_key:
        try:
            config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite", response_schema=StudyPlanSchema, capabilities=FAST_CAPS)
            prompt = f"Generate a detailed step-by-step study plan (Giáo án) based on the following roadmap/text: '{text_to_process}'. Include specific learning activities and estimated time. Format as markdown."
            data = await agent_run(config, prompt)
            if data and data.get("markdown_content"):
                if request.project_id or request.document_id:
                    query = db.query(models.Artifact).filter(models.Artifact.type == "studyplan")
                    if request.document_id:
                        query = query.filter(models.Artifact.document_id == request.document_id)
                    elif request.project_id:
                        query = query.filter(models.Artifact.project_id == request.project_id)
                    query.delete()
                    db.commit()
                    crud.create_artifact(db, project_id=request.project_id, type="studyplan", title=data.get("title", "Giáo Án"), content=json.dumps(data), document_id=request.document_id)
                return data
        except Exception as e:
            print(f"[study_plan] AI failed, falling back to offline: {e}")
            
    # Fallback Offline Engine
    from nlp.concept_map import generate_offline_study_plan
    data = generate_offline_study_plan(text_to_process)
    if request.project_id or request.document_id:
        query = db.query(models.Artifact).filter(models.Artifact.type == "studyplan")
        if request.document_id:
            query = query.filter(models.Artifact.document_id == request.document_id)
        elif request.project_id:
            query = query.filter(models.Artifact.project_id == request.project_id)
        query.delete()
        db.commit()
        crud.create_artifact(db, project_id=request.project_id, type="studyplan", title=data.get("title", "Giáo Án"), content=json.dumps(data), document_id=request.document_id)
    return data

@app.post("/api/quizzes/submit")
async def submit_quiz(request: QuizSubmitRequest, db: Session = Depends(get_db)):
    try:
        doc = db.query(models.Document).filter(models.Document.id == request.document_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        score_record = crud.create_quiz_score(
            db, 
            document_id=request.document_id, 
            score=request.score, 
            total_questions=request.total_questions
        )
        return {"success": True, "score_id": score_record.id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/documents/{document_id}/progress")
async def get_document_progress(document_id: int, db: Session = Depends(get_db)):
    try:
        from nlp.readability import score_difficulty
        
        doc = db.query(models.Document).filter(models.Document.id == document_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
            
        readability = score_difficulty(doc.content)
        scores = crud.get_quiz_scores(db, document_id)
        
        total_quizzes = len(scores)
        average_score = 0
        latest_score = None
        
        if total_quizzes > 0:
            total_points = sum([s.score for s in scores])
            total_max = sum([s.total_questions for s in scores])
            if total_max > 0:
                average_score = (total_points / total_max) * 100
            
            latest = scores[-1]
            if latest.total_questions > 0:
                latest_score = (latest.score / latest.total_questions) * 100
                
        return {
            "readability": readability,
            "average_quiz_score": round(average_score, 1),
            "latest_score": round(latest_score, 1) if latest_score is not None else None,
            "total_quizzes_taken": total_quizzes
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate_notes")
async def generate_notes(request: TopicRequest):
    current_key = get_api_key(request.api_key)
    if current_key:
        try:
            config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite", response_schema=NoteSchema, capabilities=FAST_CAPS)
            prompt = (
                "Generate high-quality, structured study notes (title, summary, and sections with "
                "bullet_points) based on the following text, in the SAME language as the text.\n"
                f"Text: {request.topic_or_text}"
            )
            data = await agent_run(config, prompt)
            if data and data.get("summary"):
                return data
        except Exception as e:
            print(f"[notes] AI failed, falling back to offline: {e}")
    from nlp.notes import extract_notes
    return extract_notes(request.topic_or_text)

@app.get("/api/flashcards/due")
async def get_due_flashcards(project_id: int | None = None, document_id: int | None = None, db: Session = Depends(get_db)):
    cards = crud.get_due_flashcards(db, project_id=project_id, document_id=document_id)
    return {"flashcards": [
        {
            "id": c.id, "front": c.front, "back": c.back,
            "interval": c.interval, "ease": c.ease, 
            "repetitions": c.repetitions, "due_date": c.due_date.isoformat(),
            "project_id": c.project_id, "document_id": c.document_id
        } for c in cards
    ]}

@app.post("/api/flashcards/review")
async def review_flashcard(request: ReviewRequest, db: Session = Depends(get_db)):
    from nlp.spaced_repetition import CardState, sm2_update
    state = CardState(
        interval=request.interval,
        ease=request.ease,
        repetitions=request.repetitions,
        due=request.due
    )
    new_state = sm2_update(state, request.quality)
    
    # If request has an ID, we update the DB
    # We need to add card_id to ReviewRequest
    if hasattr(request, "id") and getattr(request, "id") is not None:
        crud.update_flashcard_sm2(
            db, request.id, new_state.interval, new_state.ease, 
            new_state.repetitions, new_state.due
        )
    return new_state.model_dump()

@app.post("/api/score_difficulty")
async def score_text_difficulty(request: TopicRequest):
    from nlp.readability import score_difficulty
    return score_difficulty(request.topic_or_text)

@app.post("/api/generate_map")
async def generate_concept_map_endpoint(request: TopicRequest):
    current_key = get_api_key(request.api_key)
    if current_key:
        try:
            config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite", response_schema=ConceptMapSchema, capabilities=FAST_CAPS)
            prompt = f"Analyze this educational text and extract the key concepts. Generate a concept map with up to 10 nodes and their relationships. Return structured JSON with 'nodes' (id, label, definition, formula) and 'edges' (source, target, weight). Nodes must be meaningful concepts in the SAME language as the text, NOT random fragments. Include a concise definition for each concept, and a formula if applicable (especially for chemistry/math/physics). Text: '{request.topic_or_text[:5000]}'"
            data = await agent_run(config, prompt)
            if data and data.get("nodes"):
                return data
        except Exception as e:
            print(f"[map] AI failed, falling back to offline: {e}")
    from nlp.concept_map import generate_concept_map
    return generate_concept_map(request.topic_or_text)

if __name__ == "__main__":
    import uvicorn
    import sys
    
    is_frozen = getattr(sys, 'frozen', False)
    if is_frozen:
        uvicorn.run(app, host="127.0.0.1", port=8000, workers=1)
    else:
        uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)

@app.post("/api/suggestions")
async def generate_suggestions(request: TopicRequest, x_api_key: str | None = Header(default=None)):
    current_key = get_api_key(x_api_key)
    
    if current_key:
        try:
            config = LocalAgentConfig(api_key=current_key, response_schema=SuggestionSchema)
            prompt = f"""Analyze the following text corpus from a user's document library. Suggest 3 things: a topic for a structured learning path, a specific topic for a multiple-choice quiz, and a topic for a set of flashcards. Keep the suggestions concise.

Corpus:
{request.topic_or_text[:5000]}"""
            async with Agent(config) as agent:
                response = await agent.chat(prompt)
                data = await response.structured_output()
                if data:
                    return data
        except Exception as e:
            print(f"[suggestions] AI failed, falling back to offline: {e}")

    # Fallback Offline Engine
    from nlp.concept_map import generate_offline_suggestions
    return generate_offline_suggestions(request.topic_or_text)


# ═══════════════════════════════════════════════════════════════════════════════
# Studio: exams with full config, exam-room documents, sharing (drafted via 9Router)
# ═══════════════════════════════════════════════════════════════════════════════
from pydantic import BaseModel as _BaseModel
import urllib.request
import tempfile

# ── Pydantic models ──────────────────────────────────────────────────────────

class ExamSource(_BaseModel):
    type: str  # "text"|"document"|"drive"
    value: str
    name: str | None = None

class ExamAddQuestionsRequest(_BaseModel):
    count: int = 5
    api_key: str | None = None

class ExamGenRequest(_BaseModel):
    title: str = "Đề thi ôn tập"
    description: str = ""
    duration_minutes: int = 45
    num_questions: int = 10
    difficulty: str = "Trung bình"
    language: str = "Tiếng Việt"
    question_types: list[str] = ["mcq"]
    with_explanation: bool = True
    sources: list[ExamSource] = []
    api_key: str | None = None
    project_id: int | None = None
    document_id: int | None = None
    owner_email: str | None = None

class ExamDocGenRequest(_BaseModel):
    title: str = "Tài liệu ôn thi"
    description: str = ""
    objective: str = ""
    length: str = "Độ dài vừa phải"
    sources: list[ExamSource] = []
    api_key: str | None = None
    owner_email: str | None = None

class ShareRequest(_BaseModel):
    artifact_id: int | None = None
    drive_file_id: str | None = None
    drive_file_name: str | None = None
    owner_email: str
    to_email: str

# ── Writing practice (Vietnamese -> English translation, modeled on datpmt) ──
class WritingLessonRequest(_BaseModel):
    mode: str = "sentence"        # "sentence" | "paragraph" | "ielts"
    level: str = "beginner"       # "beginner" | "intermediate" | "advanced"
    category: str = ""            # sentence category or paragraph content type
    task: str = "task2"           # IELTS only: "task1" | "task2"
    topic: str = ""
    num_sentences: int = 10
    project_id: int | None = None
    owner_email: str | None = None
    api_key: str | None = None
    question_type: str | None = ""

class TeachBackPromptSchema(_BaseModel):
    concept: str
    key_points: list[str] = []
    question_vi: str

class TeachBackMisconception(_BaseModel):
    claim: str
    correction: str

class TeachBackEvalSchema(_BaseModel):
    understanding_score: int
    covered: list[str] = []
    gaps: list[str] = []
    misconceptions: list[TeachBackMisconception] = []
    followup_question: str
    feedback_vi: str

from pydantic import BaseModel
class TeachBackStartRequest(BaseModel):
    topic: str
    context: str | None = ""
    level: str | None = ""
    document_id: int | None = None
    project_id: int | None = None
    api_key: str | None = None

class TeachBackEvalRequest(BaseModel):
    concept: str
    key_points: list[str] = []
    explanation: str
    context: str | None = ""
    api_key: str | None = None


class WritingGradeRequest(_BaseModel):
    vi: str = ""
    reference_en: str = ""
    user_en: str = ""
    level: str = ""
    api_key: str | None = None

class WritingIeltsRequest(_BaseModel):
    prompt: str = ""
    essay: str = ""
    task: str = "task2"
    api_key: str | None = None
    question_type: str | None = ""

class WritingWordRequest(_BaseModel):
    word: str = ""
    sentence_en: str = ""
    api_key: str | None = None

class WritingAttemptRequest(_BaseModel):
    owner_email: str | None = None
    mode: str = "sentence"
    level: str = "beginner"
    category: str = ""
    task: str = ""
    score: float = 0.0          # sentence/paragraph: avg /10; ielts: band /9
    num_items: int = 0

class VocabAddRequest(_BaseModel):
    word: str
    meaning_vi: str | None = ""
    ipa: str | None = ""
    part_of_speech: str | None = ""
    example_en: str | None = ""
    example_vi: str | None = ""
    owner_email: str | None = None
    api_key: str | None = None

class VocabGradeRequest(_BaseModel):
    id: int
    correct: bool
    owner_email: str | None = None

# AI response schemas for writing
class WritingSentenceSchema(_BaseModel):
    vi: str
    reference_en: str
    hint: str | None = ""

class WritingLessonSchema(_BaseModel):
    mode: str
    level: str
    category: str | None = ""
    sentences: list[WritingSentenceSchema] = []
    ielts_prompt: str | None = ""
    min_words: int | None = 0
    question_type: str | None = ""
    structure_hint: list[str] = []

class WritingWordSchema(_BaseModel):
    word: str
    ipa: str | None = ""
    part_of_speech: str | None = ""
    meaning_vi: str
    example_en: str | None = ""
    example_vi: str | None = ""

class WritingErrorSchema(_BaseModel):
    type: str
    vi_explanation: str

class WritingGradeSchema(_BaseModel):
    score: float
    corrected: str
    errors: list[WritingErrorSchema] = []
    tip_vi: str
    accuracy: int = 0
    is_good: bool = False
    suggestion: str | None = ""
    improvements: list[str] = []

class WritingVocabBankRequest(BaseModel):
    text: str
    api_key: str | None = None

class WritingVocabBankSchema(_BaseModel):
    words: list[WritingWordSchema] = []

class WritingBandDescriptors(_BaseModel):
    task_response: str | None = ""
    coherence: str | None = ""
    lexical: str | None = ""
    grammar: str | None = ""

class WritingAnnotation(_BaseModel):
    original: str
    issue_vi: str
    suggestion: str

class WritingVocabItem(_BaseModel):
    basic: str
    better: str
    note_vi: str | None = ""

class WritingStructureSchema(_BaseModel):
    intro: str | None = ""
    body: str | None = ""
    conclusion: str | None = ""

class WritingIeltsSchema(_BaseModel):
    band: float
    task_response: float
    coherence: float
    lexical: float
    grammar: float
    feedback_vi: str
    improved_intro: str
    descriptors: WritingBandDescriptors | None = None
    annotations: list[WritingAnnotation] = []
    vocabulary: list[WritingVocabItem] = []
    model_answer: str | None = ""
    structure: WritingStructureSchema | None = None
    question_type: str | None = ""

# AI response schemas
class ExamOption(_BaseModel):
    id: str
    text: str

class ExamQuestionSchema(_BaseModel):
    type: str
    question: str
    options: list[ExamOption] | None = None
    correct_option_id: str | None = None
    answer: str | None = None
    explanation: str | None = None
    source_ref: str | None = None

class ExamSchema(_BaseModel):
    title: str
    questions: list[ExamQuestionSchema]

class ExamDocSchema(_BaseModel):
    title: str
    markdown_content: str


# ── Helpers ──────────────────────────────────────────────────────────────────

def _fetch_drive_text(file_id: str) -> str:
    try:
        key = os.getenv("VITE_GOOGLE_DRIVE_API_KEY", "")
        # macOS system Python lacks CA certs (CERTIFICATE_VERIFY_FAILED) — use certifi's bundle.
        import ssl, certifi
        ctx = ssl.create_default_context(cafile=certifi.where())
        # The public library is built from Drive shortcuts; alt=media 403s on a shortcut id,
        # so resolve to the target file first.
        meta_url = f"https://www.googleapis.com/drive/v3/files/{file_id}?fields=id,shortcutDetails&key={key}"
        try:
            with urllib.request.urlopen(meta_url, timeout=30, context=ctx) as resp:
                meta = json.loads(resp.read().decode())
            file_id = meta.get("shortcutDetails", {}).get("targetId") or file_id
        except Exception:
            pass
        url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media&key={key}"
        with urllib.request.urlopen(url, timeout=60, context=ctx) as resp:
            data = resp.read()
        if data[:4] == b"%PDF":
            import pdfplumber
            tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            try:
                tmp.write(data)
                tmp.close()
                with pdfplumber.open(tmp.name) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            finally:
                os.unlink(tmp.name)
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _assemble_source_text(db, sources: list, document_id=None) -> str:
    parts = []
    for src in sources:
        try:
            if src.type == "text":
                parts.append(src.value)
            elif src.type == "document":
                parts.append(get_document_text(db, int(src.value)))
            elif src.type == "drive":
                parts.append(_fetch_drive_text(src.value))
        except Exception:
            pass
    if document_id:
        try:
            parts.append(get_document_text(db, document_id))
        except Exception:
            pass
    return "\n\n".join(filter(None, parts))[:60000]

def _assemble_source_text_paged(db, sources: list, document_id=None) -> str:
    """Like _assemble_source_text but, for `document` and top-level `document_id`
    sources that have per-page text (Document.pages_data JSON array), interleaves
    a `[Trang N]` marker before each page so the model can cite pages. Falls back
    to the plain text for text/drive sources. Capped at 60000 chars."""
    import json
    parts = []
    
    def process_doc(doc_id):
        try:
            doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
            if doc and doc.pages_data:
                pages = json.loads(doc.pages_data)
                if isinstance(pages, list) and pages:
                    return "\n\n".join([f"[Trang {i+1}]\n{page_text}" for i, page_text in enumerate(pages)])
        except Exception:
            pass
        return None

    for src in sources:
        try:
            if src.type == "text":
                parts.append(src.value)
            elif src.type == "document":
                doc_id = int(src.value)
                paged_text = process_doc(doc_id)
                if paged_text:
                    parts.append(paged_text)
                else:
                    parts.append(get_document_text(db, doc_id))
            elif src.type == "drive":
                parts.append(_fetch_drive_text(src.value))
        except Exception:
            pass
    if document_id:
        try:
            paged_text = process_doc(document_id)
            if paged_text:
                parts.append(paged_text)
            else:
                parts.append(get_document_text(db, document_id))
        except Exception:
            pass
    return "\n\n".join(filter(None, parts))[:60000]


# ── Endpoints ────────────────────────────────────────────────────────────────

def _normalize_exam_questions(questions):
    """Repair AI/offline exam output so the taker can grade + highlight the right answer.
    The model is inconsistent: it sometimes puts the answer in `answer` ("A"/text) instead of
    `correct_option_id`, emits option ids with trailing commas ("A,"), or leaves
    correct_option_id null. Normalize option ids and always resolve a `correct_option_id`
    that matches one of the option ids. Best-effort; leaves the question untouched on failure."""
    def _clean(x):
        return x.strip().strip(",").strip() if isinstance(x, str) else x
    for q in (questions or []):
        try:
            opts = q.get("options")
            if not opts:
                continue
            for o in opts:
                if isinstance(o.get("id"), str):
                    o["id"] = _clean(o["id"])
            ids = [o.get("id") for o in opts]
            cid = _clean(q.get("correct_option_id"))
            if cid in ids:
                q["correct_option_id"] = cid
                continue
            ans = _clean(q.get("answer"))
            chosen = None
            if isinstance(ans, str) and ans:
                if ans in ids:
                    chosen = ans
                if chosen is None:  # answer given as the option's text
                    for o in opts:
                        if (o.get("text") or "").strip().lower() == ans.lower():
                            chosen = o.get("id"); break
                if chosen is None and len(ans) == 1 and ans.upper().isalpha():
                    idx = ord(ans.upper()) - ord("A")  # bare letter → index
                    if 0 <= idx < len(opts):
                        chosen = opts[idx].get("id")
            if chosen is not None:
                q["correct_option_id"] = chosen
        except Exception:
            pass
    return questions


@app.post("/api/exams/generate")
async def generate_exam(req: ExamGenRequest, db=Depends(get_db)):
    from nlp.exam import extract_exam

    text = _assemble_source_text_paged(db, req.sources, document_id=req.document_id)
    if not text.strip():
        text = f"{req.title}. {req.description}"

    data = None
    current_key = get_api_key(req.api_key)
    if current_key:
        try:
            prompt_text = text[:30000]
            type_names = ", ".join(req.question_types)
            expl_instruction = "Có giải thích cho mỗi câu." if req.with_explanation else "Không cần giải thích."
            if req.with_explanation:
                expl_instruction += " Nội dung nguồn được đánh dấu ranh giới trang bằng [Trang N]. Với mỗi câu hỏi, phần giải thích phải bám sát nội dung nguồn và trích dẫn trang tương ứng; đặt số trang vào trường source_ref dưới dạng 'Trang N'."
            mcq_rule = (
                " Với câu trắc nghiệm: cung cấp đúng 4 phương án; KHÔNG thêm phương án "
                "'Đáp án khác'/'Other'. Đáp án đúng PHẢI là một trong các phương án, và phải "
                "đặt id của phương án đúng vào trường correct_option_id (khớp chính xác với "
                "id của phương án đó). Kiểm tra lại phép tính để correct_option_id nhất quán "
                "với phần giải thích."
            )
            prompt = (
                f"Tạo đề thi gồm {req.num_questions} câu hỏi ({type_names}), "
                f"độ khó: {req.difficulty}, ngôn ngữ: {req.language}. {expl_instruction}{mcq_rule}\n\n"
                f"Nội dung:\n{prompt_text}"
            )
            config = LocalAgentConfig(
                api_key=current_key,
                model="gemini-3.1-flash-lite",
                response_schema=ExamSchema,
                capabilities=FAST_CAPS,
            )
            result = await agent_run(config, prompt, structured=True, timeout=30)
            if result and result.get("questions"):
                data = result
        except Exception as e:
            print(f"[exam AI failed] {e}")

    if not data:
        data = extract_exam(
            text,
            num_questions=req.num_questions,
            types=req.question_types,
            with_explanation=req.with_explanation,
        )

    data["questions"] = _normalize_exam_questions(data.get("questions", []))

    cfg = {
        "duration_minutes": req.duration_minutes,
        "difficulty": req.difficulty,
        "language": req.language,
        "question_types": req.question_types,
        "with_explanation": req.with_explanation,
        "description": req.description,
    }

    if req.project_id or req.document_id:
        query = db.query(models.Artifact).filter(models.Artifact.type == "exam")
        if req.document_id:
            query = query.filter(models.Artifact.document_id == req.document_id)
        elif req.project_id:
            query = query.filter(models.Artifact.project_id == req.project_id)
        query.delete()
        db.commit()

    artifact = crud.create_artifact(
        db,
        project_id=req.project_id,
        type="exam",
        title=req.title or data.get("title", "Đề thi"),
        content=json.dumps({"config": cfg, "questions": data["questions"]}, ensure_ascii=False),
        document_id=req.document_id,
        owner_email=req.owner_email,
    )
    return {"id": artifact.id, "title": artifact.title, "config": cfg, "questions": data["questions"]}


@app.post("/api/exams/{artifact_id}/questions")
async def add_exam_questions(artifact_id: int, req: ExamAddQuestionsRequest, db=Depends(get_db)):
    from nlp.exam import extract_exam
    artifact = db.query(models.Artifact).filter(
        models.Artifact.id == artifact_id, models.Artifact.type == "exam").first()
    if not artifact:
        raise HTTPException(status_code=404, detail="Đề thi không tìm thấy")
    try:
        content = json.loads(artifact.content)
    except Exception:
        content = {}
    cfg = content.get("config", {})
    existing = content.get("questions", [])
    count = max(1, min(int(req.count or 5), 20))
    q_types = cfg.get("question_types") or ["mcq"]
    with_expl = cfg.get("with_explanation", True)

    # Source text: prefer the exam's linked document (paged, so source_ref works);
    # otherwise extend from the existing question texts so the AI stays on-topic.
    text = ""
    if artifact.document_id:
        text = _assemble_source_text_paged(db, [], document_id=artifact.document_id)
    if not text.strip():
        text = "\n".join((q.get("question") or "") for q in existing)

    new_qs = None
    current_key = get_api_key(req.api_key)
    if current_key and text.strip():
        try:
            type_names = ", ".join(q_types)
            avoid = " / ".join((q.get("question") or "")[:80] for q in existing[:12])
            expl_instruction = ("Có giải thích cho mỗi câu. Nội dung nguồn được đánh dấu "
                "ranh giới trang bằng [Trang N]; mỗi giải thích phải trích dẫn trang tương "
                "ứng và đặt số trang vào trường source_ref dạng 'Trang N'.") if with_expl else "Không cần giải thích."
            prompt = (
                f"Tạo thêm {count} câu hỏi MỚI ({type_names}) cho đề thi, KHÁC hoàn toàn với "
                f"các câu đã có sau đây: {avoid}. {expl_instruction}\n\nNội dung:\n{text[:30000]}"
            )
            config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite",
                                      response_schema=ExamSchema, capabilities=FAST_CAPS)
            result = await agent_run(config, prompt, structured=True, timeout=30)
            if result and result.get("questions"):
                new_qs = result["questions"][:count]
        except Exception as e:
            print(f"[exam add AI failed] {e}")

    if not new_qs:
        data = extract_exam(text, num_questions=count, types=q_types, with_explanation=with_expl)
        new_qs = data.get("questions", [])

    merged = _normalize_exam_questions(existing + (new_qs or []))
    content["questions"] = merged
    artifact.content = json.dumps(content, ensure_ascii=False)
    db.commit()
    return {"id": artifact.id, "title": artifact.title, "config": cfg,
            "questions": merged, "added": len(new_qs or [])}


@app.post("/api/examdocs/generate")
async def generate_examdoc(req: ExamDocGenRequest, db=Depends(get_db)):
    from nlp.concept_map import generate_offline_exam_prep

    text = _assemble_source_text(db, req.sources)
    if not text.strip():
        text = f"{req.title}. {req.description}"

    data = None
    current_key = get_api_key(req.api_key)
    if current_key:
        try:
            prompt_text = text[:30000]
            prompt = (
                f"Tạo tài liệu phòng thi (cram sheet) toàn diện bằng tiếng Việt. "
                f"Mục tiêu: {req.objective or 'ôn tập tổng quát'}. "
                f"Độ dài mong muốn: {req.length}. "
                f"Mô tả thêm: {req.description}.\n\n"
                f"Trình bày dạng Markdown, có tiêu đề, gạch đầu dòng, bảng nếu cần. "
                f"Tóm tắt các điểm then chốt, công thức, định nghĩa quan trọng.\n\n"
                f"Nội dung nguồn:\n{prompt_text}"
            )
            config = LocalAgentConfig(
                api_key=current_key,
                model="gemini-3.1-flash-lite",
                response_schema=ExamDocSchema,
                capabilities=FAST_CAPS,
            )
            result = await agent_run(config, prompt, structured=True, timeout=30)
            if result and result.get("markdown_content"):
                data = result
        except Exception as e:
            print(f"[examdoc AI failed] {e}")

    if not data:
        data = generate_offline_exam_prep(text)

    cfg = {"objective": req.objective, "length": req.length, "description": req.description}
    artifact = crud.create_artifact(
        db,
        type="examdoc",
        title=req.title or data.get("title", "Tài liệu ôn thi"),
        content=json.dumps({"config": cfg, "markdown_content": data["markdown_content"]}, ensure_ascii=False),
        owner_email=req.owner_email,
    )
    return {
        "id": artifact.id,
        "title": artifact.title,
        "markdown_content": data["markdown_content"],
        "config": cfg,
    }


@app.get("/api/artifacts/{artifact_id}")
def get_artifact(artifact_id: int, db=Depends(get_db)):
    artifact = db.query(models.Artifact).filter(models.Artifact.id == artifact_id).first()
    if not artifact:
        raise HTTPException(status_code=404, detail="Artifact không tìm thấy")
    try:
        content = json.loads(artifact.content)
    except Exception:
        content = artifact.content
    # Repair legacy/AI exam payloads on read (correct_option_id missing / option ids with
    # stray commas) so older exams grade + highlight correctly without regenerating.
    if artifact.type == "exam" and isinstance(content, dict) and content.get("questions"):
        content["questions"] = _normalize_exam_questions(content["questions"])
    return {
        "id": artifact.id,
        "type": artifact.type,
        "title": artifact.title,
        "content": content,
        "document_id": artifact.document_id,
        "project_id": artifact.project_id,
        "created_at": artifact.created_at.isoformat() if artifact.created_at else None,
        "owner_email": artifact.owner_email,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# Writing practice: VN -> EN translation lessons, sentence grading, IELTS scoring.
# Every AI path falls back to the deterministic nlp.writing engine (never 500).
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/writing/lessons/generate")
async def writing_generate_lesson(req: WritingLessonRequest, db=Depends(get_db)):
    from nlp import writing as writing_engine

    data = None
    current_key = get_api_key(req.api_key)
    if current_key:
        try:
            if req.mode == "ielts":
                task_label = "Task 1 (biểu đồ)" if req.task == "task1" else "Task 2 (bài luận)"
                if req.task == "task2":
                    ask = (
                        f"Tạo một bài luận IELTS Task 2 hoàn chỉnh đạt band 8 trả lời cho một đề bài ngẫu nhiên "
                        f"thuộc loại '{req.question_type or 'bất kỳ'}'. Sau đó, tách bài luận thành từng câu. "
                        "Với mỗi câu, cung cấp bản dịch tiếng Việt tự nhiên ('vi') và câu tiếng Anh gốc ('reference_en'), hint để rỗng. "
                        "Trả về thêm trường 'question_type' và 'structure_hint' (dàn ý gợi ý 4 mục bằng tiếng Việt)."
                    )
                else:
                    ask = f"Tạo một đề bài IELTS Writing {task_label} ngẫu nhiên tiếng Anh."
                prompt = (
                    ask + "\n\nTrả về JSON đúng schema: mode, level, category, "
                    "sentences[{vi, reference_en, hint}], ielts_prompt, min_words (250 cho task2), question_type, structure_hint."
                )
            elif req.mode == "paragraph":
                ask = (
                    f"Tạo một đoạn văn tiếng Việt mạch lạc thuộc thể loại '{req.category or 'Essays'}' "
                    f"ở trình độ {req.level}, rồi tách thành {req.num_sentences} câu. "
                    "Với mỗi câu, cung cấp bản dịch tiếng Anh chuẩn."
                )
                prompt = (
                    ask + "\n\nTrả về JSON đúng schema: mode, level, category, "
                    "sentences[{vi, reference_en, hint}], ielts_prompt, min_words."
                )
            else:
                ask = (
                    f"Tạo {req.num_sentences} câu tiếng Việt thuộc chủ đề '{req.category or 'đời sống hằng ngày'}' "
                    f"ở trình độ {req.level} để luyện dịch sang tiếng Anh. "
                    "Với mỗi câu, cung cấp bản dịch tiếng Anh chuẩn và một gợi ý ngắn (hint) bằng tiếng Việt."
                )
                prompt = (
                    ask + "\n\nTrả về JSON đúng schema: mode, level, category, "
                    "sentences[{vi, reference_en, hint}], ielts_prompt, min_words."
                )
            config = LocalAgentConfig(
                api_key=current_key, model="gemini-3.1-flash-lite",
                response_schema=WritingLessonSchema, capabilities=FAST_CAPS,
            )
            result = await agent_run(config, prompt, structured=True, timeout=30)
            
            # Additional fallback checking for IELTS
            if result and (result.get("sentences") or req.mode == "ielts"):
                result.setdefault("mode", req.mode)
                result.setdefault("level", req.level)
                result.setdefault("category", req.category)
                if req.mode == "ielts":
                    if req.task == "task2" and not result.get("sentences"):
                        # Fallback for omitted sentences or structure
                        fallback_data = writing_engine.generate_lesson(
                            mode=req.mode, level=req.level, category=req.category,
                            topic=req.topic, num_sentences=req.num_sentences, task=req.task,
                            question_type=req.question_type,
                        )
                        result["sentences"] = fallback_data.get("sentences", [])
                        if not result.get("question_type"):
                            result["question_type"] = fallback_data.get("question_type", "")
                        if not result.get("structure_hint"):
                            result["structure_hint"] = fallback_data.get("structure_hint", [])
                        if not result.get("ielts_prompt"):
                            result["ielts_prompt"] = fallback_data.get("ielts_prompt", "")
                    else:
                        result.setdefault("sentences", [])
                    result.setdefault("min_words", 150 if req.task == "task1" else 250)
                else:
                    result.setdefault("ielts_prompt", "")
                    result.setdefault("min_words", 0)
                data = result
        except Exception as e:
            print(f"[writing lesson AI failed] {e}")

    if not data:
        data = writing_engine.generate_lesson(
            mode=req.mode, level=req.level, category=req.category,
            topic=req.topic, num_sentences=req.num_sentences, task=req.task,
            question_type=req.question_type,
        )

    title_map = {"sentence": "Luyện câu", "paragraph": "Luyện đoạn văn", "ielts": "Luyện IELTS Writing"}
    title = f"{title_map.get(req.mode, 'Luyện viết')} · {req.category or req.level}"
    artifact = crud.create_artifact(
        db, project_id=req.project_id, type="writing_lesson", title=title,
        content=json.dumps(data, ensure_ascii=False),
        owner_email=req.owner_email,
    )
    return {"id": artifact.id, "title": title, **data}


@app.post("/api/writing/grade")
async def writing_grade(req: WritingGradeRequest):
    from nlp import writing as writing_engine

    current_key = get_api_key(req.api_key)
    if current_key and req.user_en.strip():
        try:
            prompt = (
                "Bạn là huấn luyện viên dịch thuật (translation coach). Chấm điểm bản dịch của học viên.\n"
                f"Câu tiếng Việt: {req.vi}\n"
                f"Đáp án tham khảo: {req.reference_en}\n"
                f"Bản dịch của học viên: {req.user_en}\n"
                "Trả về JSON khớp với WritingGradeSchema gồm: "
                "score (0-10), accuracy (0-100, độ sát nghĩa và ngữ pháp so với câu dịch chuẩn), "
                "is_good (true nếu dịch chính xác và tự nhiên), "
                "corrected (phiên bản đã sửa/tốt hơn), "
                "suggestion (câu đã sửa/câu tham khảo nếu chưa tốt, ngược lại rỗng), "
                "errors[{type, vi_explanation}], "
                "improvements (2-4 gạch đầu dòng ngắn gọn bằng tiếng Việt chỉ ra điểm cần sửa: thiếu ý, sai cấu trúc, sai thì...; "
                "BỌC các từ/cụm tiếng Anh quan trọng trong dấu **...** để làm nổi bật, ví dụ: 'nên dùng **mixed chart** thay vì **chaotic graph**'), "
                "tip_vi (1 câu khen ngợi nếu tốt, hoặc 1 câu khuyên nhủ/nhắc nhở từ khoá cần dùng nếu chưa tốt; cũng bọc từ khoá tiếng Anh trong **...**)."
            )
            config = LocalAgentConfig(
                api_key=current_key, model="gemini-3.1-flash-lite",
                response_schema=WritingGradeSchema, capabilities=FAST_CAPS,
            )
            result = await agent_run(config, prompt, structured=True, timeout=25)
            if result and "score" in result:
                result.setdefault("errors", [])
                result.setdefault("accuracy", int(result["score"] * 10))
                result.setdefault("is_good", result["score"] >= 8.0)
                result.setdefault("suggestion", req.reference_en if not result["is_good"] else "")
                result.setdefault("improvements", [])
                result["reference_en"] = req.reference_en
                return result
        except Exception as e:
            print(f"[writing grade AI failed] {e}")

    return writing_engine.grade_sentence(req.vi, req.reference_en, req.user_en, req.level)


@app.post("/api/writing/ielts/grade")
async def writing_grade_ielts(req: WritingIeltsRequest):
    from nlp import writing as writing_engine

    task_label = "Task 1 (mô tả biểu đồ/quy trình, báo cáo trang trọng, tối thiểu 150 từ)" \
        if req.task == "task1" else "Task 2 (bài luận nêu quan điểm, tối thiểu 250 từ)"
    current_key = get_api_key(req.api_key)
    if current_key and req.essay.strip():
        try:
            prompt = (
                f"Bạn là giám khảo IELTS khó tính nhưng luôn khuyến khích học viên. Chấm bài Writing {task_label} dưới đây IN VIETNAMESE.\n"
                f"Đề bài: {req.prompt}\n\nBài làm:\n{req.essay}\n\n"
                "Cho band tổng và 4 tiêu chí (task_response, coherence, lexical, grammar) theo thang 0-9 (bước 0.5). "
                "Viết feedback_vi bằng tiếng Việt (vài gạch đầu dòng), "
                "improved_intro là đoạn mở bài được cải thiện (tiếng Anh), và descriptors là mô tả ngắn cho band đạt được. "
                "Cung cấp annotations (4-8 câu cần sửa từ bài làm), vocabulary (5-8 từ nên nâng cấp), "
                "model_answer (bài mẫu band 8 tiếng Anh), structure (đánh giá cấu trúc intro, body, conclusion bằng tiếng Việt). "
                f"question_type là '{req.question_type or ''}'."
            )
            config = LocalAgentConfig(
                api_key=current_key, model="gemini-3.1-flash-lite",
                response_schema=WritingIeltsSchema, capabilities=FAST_CAPS,
            )
            result = await agent_run(config, prompt, structured=True, timeout=35)
            if result and "band" in result:
                result.setdefault("task", req.task)
                result.setdefault("question_type", req.question_type)
                return result
        except Exception as e:
            print(f"[writing ielts AI failed] {e}")

    return writing_engine.grade_ielts(req.prompt, req.essay, task=req.task, question_type=req.question_type)


@app.post("/api/writing/word")
async def writing_word_lookup(req: WritingWordRequest):
    """Click-a-word dictionary lookup — the 'interactive translation' feature."""
    from nlp import writing as writing_engine

    current_key = get_api_key(req.api_key)
    if current_key and req.word.strip():
        try:
            prompt = (
                f"Tra cứu từ tiếng Anh '{req.word}'"
                + (f" trong câu: \"{req.sentence_en}\"." if req.sentence_en else ".")
                + " Trả về: word (từ gốc), ipa (phiên âm IPA), part_of_speech (từ loại), "
                "meaning_vi (nghĩa tiếng Việt ngắn gọn), example_en (một câu ví dụ tiếng Anh), "
                "example_vi (bản dịch tiếng Việt của ví dụ)."
            )
            config = LocalAgentConfig(
                api_key=current_key, model="gemini-3.1-flash-lite",
                response_schema=WritingWordSchema, capabilities=FAST_CAPS,
            )
            result = await agent_run(config, prompt, structured=True, timeout=20)
            if result and result.get("meaning_vi"):
                return result
        except Exception as e:
            print(f"[writing word AI failed] {e}")

    return writing_engine.lookup_word(req.word, req.sentence_en)

@app.post("/api/writing/vocab_bank")
async def writing_vocab_bank(req: WritingVocabBankRequest):
    from nlp import writing as writing_engine
    words = None
    current_key = get_api_key(req.api_key)
    if current_key and req.text.strip():
        try:
            config = LocalAgentConfig(
                response_schema=WritingVocabBankSchema,
                model="gemini-3.1-flash-lite",
                capabilities=FAST_CAPS,
                api_key=current_key
            )
            prompt = (
                "Bạn là trợ lý từ vựng IELTS. Cho ĐOẠN VĂN tiếng Anh dưới đây, trích tối đa 12 "
                "từ/cụm từ quan trọng người học nên ghi nhớ. Với MỖI từ trả về: word; ipa (phiên "
                "âm); part_of_speech (từ loại); meaning_vi = nghĩa tiếng Việt ĐÚNG THEO NGỮ CẢNH "
                "của từ TRONG đoạn văn này — chọn đúng nét nghĩa đang được dùng, TUYỆT ĐỐI không "
                "đưa nghĩa chung chung hay nét nghĩa khác không xuất hiện ở đây; example_en = trích "
                "NGUYÊN VĂN câu trong đoạn văn có chứa từ đó.\n\nĐoạn văn:\n" + req.text
            )
            result = await agent_run(config, prompt, structured=True, timeout=20)
            if result and result.get("words"):
                words = result["words"]
        except Exception as e:
            print(f"[vocab bank AI failed] {e}")

    # Offline extraction (TextRank/keywords) picks the words; the Dictionary API then fills
    # real details so definitions show even without an AI key (a general, not contextual, sense).
    if words is None:
        words = writing_engine.vocab_bank(req.text)
    words = await _enrich_words(words, context=req.text)
    # Contextual example: prefer the actual sentence from the learner's passage that contains
    # the word, so a saved word carries the example of what they were writing (not a generic one).
    for w in (words or []):
        try:
            ex = _passage_example(w.get("word"), req.text)
            if ex:
                w["example_en"] = ex
        except Exception:
            pass
    return {"words": words}


# End of writing endpoints, append vocab endpoints here
def _vocab_status(w) -> str:
    """datpmt-style status: a saved-but-unmastered word is 'learning' (Đang học);
    once repetitions>=3 it flips to learned → 'learned' (Đã thuộc)."""
    return "learned" if (w.learned or (w.repetitions or 0) >= 3) else "learning"


@app.post("/api/vocabulary/add")
async def add_vocab(req: VocabAddRequest, db: Session = Depends(get_db)):
    # Dedup: if this owner already saved this word, return the existing entry (and its
    # status) instead of creating a duplicate — this is what wires the writing Dictionary
    # to the Vocabulary trainer (re-saving a word is a no-op that reports its status).
    from sqlalchemy import func
    wlow = (req.word or "").strip().lower()
    if wlow:
        dq = db.query(models.VocabWord).filter(func.lower(models.VocabWord.word) == wlow)
        dq = dq.filter(models.VocabWord.owner_email == req.owner_email) if req.owner_email \
            else dq.filter(models.VocabWord.owner_email.is_(None))
        existing = dq.first()
        if existing:
            return {
                "id": existing.id, "word": existing.word, "meaning_vi": existing.meaning_vi,
                "ipa": existing.ipa, "part_of_speech": existing.part_of_speech,
                "example_en": existing.example_en, "example_vi": existing.example_vi,
                "learned": existing.learned, "correct_count": existing.correct_count,
                "wrong_count": existing.wrong_count,
                "already": True, "status": _vocab_status(existing),
            }

    meaning = req.meaning_vi or ""
    ipa = req.ipa or ""
    pos = req.part_of_speech or ""
    ex_en = req.example_en or ""
    ex_vi = req.example_vi or ""

    if not meaning:
        current_key = get_api_key(req.api_key)
        if current_key:
            try:
                config = LocalAgentConfig(api_key=current_key, model="gemini-3.1-flash-lite", response_schema=WritingWordSchema, capabilities=FAST_CAPS)
                prompt = f"Provide phonetic transcription (ipa), part of speech (part_of_speech), Vietnamese meaning (meaning_vi), and one English example sentence (example_en) with Vietnamese translation (example_vi) for the word: '{req.word}'."
                data = await agent_run(config, prompt)
                if data:
                    meaning = data.get("meaning_vi", meaning)
                    ipa = data.get("ipa", ipa)
                    pos = data.get("part_of_speech", pos)
                    ex_en = data.get("example_en", ex_en)
                    ex_vi = data.get("example_vi", ex_vi)
            except Exception as e:
                print(f"[vocab] AI autofill failed, falling back: {e}")
                from nlp import writing as writing_engine
                try:
                    data = writing_engine.lookup_word(req.word)
                    if data:
                        meaning = data.get("meaning_vi", meaning)
                        ipa = data.get("ipa", ipa)
                        pos = data.get("part_of_speech", pos)
                        ex_en = data.get("example_en", ex_en)
                        ex_vi = data.get("example_vi", ex_vi)
                except Exception:
                    pass
        else:
            from nlp import writing as writing_engine
            try:
                data = writing_engine.lookup_word(req.word)
                if data:
                    meaning = data.get("meaning_vi", meaning)
                    ipa = data.get("ipa", ipa)
                    pos = data.get("part_of_speech", pos)
                    ex_en = data.get("example_en", ex_en)
                    ex_vi = data.get("example_vi", ex_vi)
            except Exception:
                pass

    new_word = models.VocabWord(
        word=req.word,
        meaning_vi=meaning,
        ipa=ipa,
        part_of_speech=pos,
        example_en=ex_en,
        example_vi=ex_vi,
        owner_email=req.owner_email,
        due_date=datetime.date.today()
    )
    db.add(new_word)
    db.commit()
    db.refresh(new_word)
    return {
        "id": new_word.id, "word": new_word.word, "meaning_vi": new_word.meaning_vi,
        "ipa": new_word.ipa, "part_of_speech": new_word.part_of_speech,
        "example_en": new_word.example_en, "example_vi": new_word.example_vi,
        "learned": new_word.learned, "correct_count": new_word.correct_count,
        "wrong_count": new_word.wrong_count,
        "already": False, "status": _vocab_status(new_word)
    }

from nlp.vocab_bank import VOCAB_BANK

@app.get("/api/vocabulary/bank")
def get_vocab_bank_topics():
    topics = list(VOCAB_BANK.keys())
    levels = ["b", "i", "a"]
    return {"topics": topics, "levels": levels}

@app.get("/api/vocabulary/bank/words")
def get_vocab_bank_words(topic: str, level: str):
    if topic in VOCAB_BANK and level in VOCAB_BANK[topic]:
        return {"words": VOCAB_BANK[topic][level]}
    return {"words": []}


class VocabGenerateRequest(BaseModel):
    prompt: str
    owner_email: str | None = None
    api_key: str | None = None

class VocabListSchema(BaseModel):
    words: list[WritingWordSchema]

@app.post("/api/vocabulary/generate")
async def generate_vocab(req: VocabGenerateRequest, db: Session = Depends(get_db)):
    current_key = get_api_key(req.api_key)
    if not current_key:
        raise HTTPException(status_code=400, detail="Vui lòng cung cấp Gemini API Key để tạo từ vựng bằng AI.")
        
    try:
        from core.agent import LocalAgentConfig, agent_run, FAST_CAPS
        config = LocalAgentConfig(
            api_key=current_key, 
            model="gemini-3.1-flash", 
            response_schema=VocabListSchema, 
            capabilities=FAST_CAPS
        )
        prompt = f"Generate a list of exactly 5 useful vocabulary words related to the topic: '{req.prompt}'. For each word, provide its phonetic transcription (ipa), part of speech (part_of_speech), Vietnamese meaning (meaning_vi), one English example sentence (example_en), and its Vietnamese translation (example_vi)."
        
        data = await agent_run(config, prompt)
        
        if not data or "words" not in data:
            return {"words": [], "added": 0}
            
        words = data["words"]
        added_count = 0
        
        # Query existing words
        query = db.query(models.VocabWord)
        if req.owner_email:
            query = query.filter(models.VocabWord.owner_email == req.owner_email)
        else:
            query = query.filter(models.VocabWord.owner_email.is_(None))
        existing_words = {w.word.lower(): w for w in query.all()}
        
        for w in words:
            w_lower = w["word"].strip().lower()
            if not w_lower: continue
            
            if w_lower not in existing_words:
                new_word = models.VocabWord(
                    word=w_lower,
                    meaning_vi=w.get("meaning_vi", ""),
                    ipa=w.get("ipa", ""),
                    part_of_speech=w.get("part_of_speech", ""),
                    example_en=w.get("example_en", ""),
                    example_vi=w.get("example_vi", ""),
                    owner_email=req.owner_email,
                    learned=0,
                    correct_count=0,
                    wrong_count=0,
                    interval=0,
                    ease=2.5,
                    repetitions=0,
                    due_date=datetime.date.today()
                )
                db.add(new_word)
                added_count += 1
                
        db.commit()
        return {"added": added_count, "words": words}
        
    except Exception as e:
        print(f"Generate vocab error: {e}")
        raise HTTPException(status_code=500, detail=str(e))



class VocabImportRequest(BaseModel):
    topic: str
    level: str
    owner_email: str | None = None

@app.post("/api/vocabulary/bank/import")
def import_vocab_bank(req: VocabImportRequest, db: Session = Depends(get_db)):
    if req.topic not in VOCAB_BANK or req.level not in VOCAB_BANK[req.topic]:
        raise HTTPException(status_code=404, detail="Topic or level not found in bank")
        
    words_to_add = VOCAB_BANK[req.topic][req.level]
    existing_words = db.query(models.VocabWord).filter(
        models.VocabWord.owner_email == req.owner_email
    ).all()
    existing_dict = {w.word.lower(): w for w in existing_words}
    
    added_count = 0
    updated_count = 0
    skipped_count = 0
    for w in words_to_add:
        w_lower = w["word"].lower()
        if w_lower not in existing_dict:
            new_word = models.VocabWord(
                word=w["word"],
                meaning_vi=w.get("meaning_vi", ""),
                ipa=w.get("ipa", ""),
                part_of_speech=w.get("part_of_speech", ""),
                example_en=w.get("example_en", ""),
                example_vi=w.get("example_vi", ""),
                owner_email=req.owner_email,
                learned=0,
                correct_count=0,
                wrong_count=0,
                interval=0,
                ease=2.5,
                repetitions=0,
                due_date=datetime.date.today()
            )
            db.add(new_word)
            added_count += 1
        else:
            existing_word = existing_dict[w_lower]
            # Update missing details if any
            if not existing_word.meaning_vi:
                existing_word.meaning_vi = w.get("meaning_vi", "")
                existing_word.ipa = existing_word.ipa or w.get("ipa", "")
                existing_word.part_of_speech = existing_word.part_of_speech or w.get("part_of_speech", "")
                existing_word.example_en = existing_word.example_en or w.get("example_en", "")
                existing_word.example_vi = existing_word.example_vi or w.get("example_vi", "")
                updated_count += 1
            else:
                skipped_count += 1
            
    db.commit()
    return {"added": added_count, "updated": updated_count, "skipped": skipped_count}

@app.get("/api/vocabulary/list")
def list_vocab(email: str | None = None, q: str | None = None, db: Session = Depends(get_db)):
    query = db.query(models.VocabWord)
    if email:
        query = query.filter(models.VocabWord.owner_email == email)
    else:
        query = query.filter(models.VocabWord.owner_email.is_(None))
    
    words = query.order_by(models.VocabWord.id.desc()).all()
    suggestions = []
    
    if q:
        q_lower = q.lower()
        
        # Generate synonyms for the query
        syns = set()
        for syn in wn.synsets(q_lower):
            for l in syn.lemmas():
                name = l.name().replace('_', ' ').lower()
                if name != q_lower:
                    syns.add(name)
        suggestions = list(syns)[:10]
        
        # Filter words matching query OR matching any synonym
        filtered_words = []
        for w in words:
            w_word = (w.word or "").lower()
            w_mean = (w.meaning_vi or "").lower()
            if q_lower in w_word or q_lower in w_mean:
                filtered_words.append(w)
            elif any(syn in w_word for syn in syns):
                filtered_words.append(w)
        words = filtered_words
        
        # If no words and no synonyms found, it might be a typo, suggest correct spellings
        if not suggestions and not words:
            if 'ALL_LEMMAS' in globals() and ALL_LEMMAS:
                suggestions = difflib.get_close_matches(q_lower, ALL_LEMMAS, n=5, cutoff=0.7)
        
    return {
        "words": [{
            "id": w.id, "word": w.word, "meaning_vi": w.meaning_vi,
            "ipa": w.ipa, "part_of_speech": w.part_of_speech,
            "example_en": w.example_en, "example_vi": w.example_vi,
            "learned": w.learned, "correct_count": w.correct_count,
            "wrong_count": w.wrong_count, "due_date": w.due_date.isoformat() if w.due_date else None
        } for w in words],
        "suggestions": suggestions
    }

def _dictionary_lookup(word: str, context: str = ""):
    """Fetch real word details from the free Dictionary API (api.dictionaryapi.dev):
    IPA, part of speech, English definition + example, then translate meaning/example to
    Vietnamese via GoogleTranslator. When `context` (the passage the word appears in) is given,
    pick the SENSE whose definition overlaps the context most — so 'media' in a communication
    passage resolves to the mass-media sense, not the anatomical one. Returns None on failure."""
    import requests, re
    try:
        from deep_translator import GoogleTranslator
        translator = GoogleTranslator(source='en', target='vi')
    except ImportError:
        translator = None

    try:
        r = requests.get(f"https://api.dictionaryapi.dev/api/v2/entries/en/{word}", timeout=5)
        if r.status_code == 200:
            json_data = r.json()
            
            # Aggregate synonyms from all objects
            all_synonyms = []
            for item in json_data:
                for m in item.get("meanings", []):
                    all_synonyms.extend(m.get("synonyms", []))
                    for d in m.get("definitions", []):
                        all_synonyms.extend(d.get("synonyms", []))
            
            # IPA from the first object.
            data = json_data[0]
            ipa = data.get("phonetic") or (data.get("phonetics") and len(data["phonetics"]) > 0 and data["phonetics"][0].get("text")) or ""

            # Collect every sense (pos + definition + example) across all objects, then pick the
            # one whose definition shares the most content words with the passage context.
            candidates = []
            for item in json_data:
                for m in item.get("meanings", []):
                    mpos = m.get("partOfSpeech", "")
                    for dfn in m.get("definitions", []):
                        candidates.append((mpos, dfn.get("definition", ""), dfn.get("example", "")))

            pos, definition, example_en, grammar_tag = "", "", "", ""
            if candidates:
                best = candidates[0]
                ctx_words = set(re.findall(r'[a-z]{4,}', (context or "").lower()))
                if ctx_words:
                    def _score(c):
                        dwords = set(re.findall(r'[a-z]{4,}', (c[1] or "").lower()))
                        return len(dwords & ctx_words)
                    scored = max(candidates, key=_score)
                    if _score(scored) > 0:
                        best = scored
                pos, definition, example_en = best[0], best[1], best[2]

            if definition:

                # Extract grammar tags robustly
                def_lower = definition.lower()
                tags = []
                
                if "plural" in def_lower and "plurality" not in def_lower:
                    tags.append("số nhiều")
                elif "singular" in def_lower and "singularity" not in def_lower:
                    tags.append("số ít")
                    
                if "uncountable" in def_lower:
                    tags.append("không đếm được")
                elif "countable" in def_lower:
                    tags.append("đếm được")
                    
                if tags:
                    grammar_tag = ", ".join(tags)
                    if pos == "noun":
                        grammar_tag = "danh từ " + grammar_tag
            
            meaning_vi = ""
            example_vi = ""
            
            if translator:
                short_meaning = ""
                long_meaning = ""
                try: short_meaning = translator.translate(word)
                except: pass
                
                if definition:
                    try: long_meaning = translator.translate(definition)
                    except: long_meaning = definition
                
                if short_meaning and long_meaning and short_meaning.lower() != word.lower():
                    meaning_vi = f"{short_meaning} - {long_meaning}"
                else:
                    meaning_vi = long_meaning or short_meaning

                if example_en:
                    try: example_vi = translator.translate(example_en)
                    except: example_vi = example_en
            else:
                meaning_vi = definition
                example_vi = example_en
                
            return {
                "word": word,
                "ipa": ipa,
                "part_of_speech": pos,
                "grammar": grammar_tag,
                "meaning_vi": meaning_vi,
                "example_en": example_en,
                "example_vi": example_vi,
                "synonyms": list(set(all_synonyms))[:8]
            }
    except Exception as e:
        print("Dictionary lookup error:", e)
    return None


@app.get("/api/vocabulary/lookup")
def lookup_dictionary(word: str):
    return _dictionary_lookup(word)


def _passage_example(word, text) -> str:
    """Return the sentence from the learner's passage that contains `word` (case-insensitive,
    whole-word preferred), so a Dictionary word carries the example of what they were actually
    writing. Empty string if the word doesn't appear verbatim (e.g. a lemma)."""
    if not word or not text:
        return ""
    import re
    wl = str(word).lower().strip()
    if not wl:
        return ""
    sentences = re.split(r'(?<=[.!?])\s+', str(text))
    # whole-word match first, then loose substring
    for pat in (r'\b' + re.escape(wl) + r'\b', re.escape(wl)):
        for s in sentences:
            if re.search(pat, s.lower()):
                s = s.strip()
                if s:
                    return s
    return ""


async def _enrich_words(words: list, context: str = "") -> list:
    """Fill in missing word details (ipa/part_of_speech/meaning_vi/example_en) from the
    Dictionary API — used so the Writing Dictionary shows real definitions even without an
    AI key. Only words lacking a Vietnamese meaning are looked up; the passage `context` steers
    the Dictionary API toward the right sense. Runs concurrently, fully best-effort."""
    import asyncio
    loop = asyncio.get_event_loop()

    async def enrich(w):
        try:
            if not isinstance(w, dict) or not w.get("word"):
                return w
            if (w.get("ipa") or "").strip():
                return w  # already enriched (the AI path fills ipa; the offline stub does not)
            detail = await loop.run_in_executor(None, _dictionary_lookup, w["word"], context)
            if detail:
                # Real details replace the offline "Bật AI (Cloud)…" placeholder.
                for k in ("ipa", "part_of_speech", "meaning_vi", "example_en", "example_vi"):
                    if detail.get(k):
                        w[k] = detail[k]
        except Exception:
            pass
        return w

    return await asyncio.gather(*[enrich(w) for w in (words or [])])


@app.delete("/api/vocabulary/{word_id}")
def delete_vocab(word_id: int, db: Session = Depends(get_db)):
    db.query(models.VocabReview).filter(models.VocabReview.vocab_id == word_id).delete()
    db.query(models.VocabWord).filter(models.VocabWord.id == word_id).delete()
    db.commit()
    return {"ok": True}

@app.get("/api/vocabulary/review")
def get_vocab_review(email: str | None = None, type: str = "new_words", limit: int = 10, db: Session = Depends(get_db)):
    limit = max(1, min(limit, 50))
    query = db.query(models.VocabWord)
    if email:
        query = query.filter(models.VocabWord.owner_email == email)
    else:
        query = query.filter(models.VocabWord.owner_email.is_(None))
        
    if type == "new_words":
        # Always fetch unlearned words, prioritize by due_date so due words come first
        query = query.filter(models.VocabWord.learned == 0)
        query = query.order_by(models.VocabWord.due_date.asc())
    else:
        query = query.filter(models.VocabWord.learned == 1)
        query = query.filter(models.VocabWord.due_date <= datetime.date.today())
        query = query.order_by(models.VocabWord.due_date.asc())
        
    words = query.limit(limit).all()
    
    if type == "new_words" and len(words) < limit:
        # Automatically pull completely new words from the global dictionary
        import random
        from nlp.vocab_bank import VOCAB_BANK
        
        # Get all existing word strings for this user
        existing_query = db.query(models.VocabWord.word)
        if email:
            existing_query = existing_query.filter(models.VocabWord.owner_email == email)
        else:
            existing_query = existing_query.filter(models.VocabWord.owner_email.is_(None))
        existing_words = {r[0].lower() for r in existing_query.all()}
        
        # Flatten all bank words
        all_bank_words = []
        for topic, levels in VOCAB_BANK.items():
            for lvl, wlist in levels.items():
                all_bank_words.extend(wlist)
                
        # Filter available words
        available_words = [w for w in all_bank_words if w["word"].lower() not in existing_words]
        
        needed = limit - len(words)
        if available_words:
            chosen = random.sample(available_words, min(needed, len(available_words)))
            for w in chosen:
                new_word = models.VocabWord(
                    word=w["word"].lower(),
                    meaning_vi=w.get("meaning_vi", ""),
                    ipa=w.get("ipa", ""),
                    part_of_speech=w.get("part_of_speech", ""),
                    example_en=w.get("example_en", ""),
                    example_vi=w.get("example_vi", ""),
                    owner_email=email,
                    learned=0,
                    correct_count=0,
                    wrong_count=0,
                    interval=0,
                    ease=2.5,
                    repetitions=0,
                    due_date=datetime.date.today()
                )
                db.add(new_word)
                words.append(new_word)
            
            db.commit()
    return {"words": [{
        "id": w.id, "word": w.word, "meaning_vi": w.meaning_vi,
        "ipa": w.ipa, "part_of_speech": w.part_of_speech,
        "example_en": w.example_en, "example_vi": w.example_vi
    } for w in words]}

@app.post("/api/vocabulary/grade")
def grade_vocab(req: VocabGradeRequest, db: Session = Depends(get_db)):
    query = db.query(models.VocabWord).filter(models.VocabWord.id == req.id)
    if req.owner_email:
        query = query.filter(models.VocabWord.owner_email == req.owner_email)
    word = query.first()
    if not word:
        raise HTTPException(status_code=404, detail="Vocab not found")
        
    if req.correct:
        word.repetitions += 1
        word.correct_count += 1
        if word.repetitions == 1:
            word.interval = 1
        elif word.repetitions == 2:
            word.interval = 2
        else:
            word.interval = round(word.interval * word.ease)
        word.ease = min(3.0, word.ease + 0.1)
        points = 10
        if word.repetitions >= 3:
            word.learned = 1
    else:
        word.wrong_count += 1
        word.repetitions = 0
        word.interval = 1
        word.ease = max(1.3, word.ease - 0.2)
        points = 2
        
    word.due_date = datetime.date.today() + datetime.timedelta(days=word.interval)
    word.last_reviewed = datetime.datetime.utcnow()
    
    review = models.VocabReview(
        owner_email=req.owner_email,
        vocab_id=req.id,
        correct=1 if req.correct else 0,
        points=points
    )
    db.add(review)
    db.commit()
    return {"id": word.id, "learned": word.learned, "points_delta": points, "repetitions": word.repetitions}

@app.get("/api/vocabulary/stats")
def get_vocab_stats(email: str | None = None, db: Session = Depends(get_db)):
    from sqlalchemy import func
    
    word_query = db.query(models.VocabWord)
    if email:
        word_query = word_query.filter(models.VocabWord.owner_email == email)
    else:
        word_query = word_query.filter(models.VocabWord.owner_email.is_(None))
        
    total = word_query.count()
    learned = word_query.filter(models.VocabWord.learned == 1).count()
    completion = round(learned / total * 100) if total > 0 else 0
    
    review_query = db.query(models.VocabReview)
    if email:
        review_query = review_query.filter(models.VocabReview.owner_email == email)
    else:
        review_query = review_query.filter(models.VocabReview.owner_email.is_(None))
        
    points = db.query(func.sum(models.VocabReview.points)).filter(
        models.VocabReview.owner_email == email if email else models.VocabReview.owner_email.is_(None)
    ).scalar() or 0
    
    # Streak calculation
    dates = db.query(func.date(models.VocabReview.created_at)).filter(
        models.VocabReview.owner_email == email if email else models.VocabReview.owner_email.is_(None)
    ).distinct().order_by(func.date(models.VocabReview.created_at).desc()).all()
    
    streak = 0
    today = datetime.date.today()
    yesterday = today - datetime.timedelta(days=1)
    
    if dates:
        # SQLite func.date() returns ISO strings ('2026-07-05'); coerce to date objects
        # so membership tests against `today`/`yesterday` (datetime.date) actually match.
        date_set = set()
        for d in dates:
            v = d[0]
            if not v:
                continue
            if isinstance(v, str):
                try:
                    v = datetime.date.fromisoformat(v[:10])
                except ValueError:
                    continue
            date_set.add(v)
        if today in date_set or yesterday in date_set:
            curr = today if today in date_set else yesterday
            while curr in date_set:
                streak += 1
                curr -= datetime.timedelta(days=1)
                
    achievements = [
        {"id": "first_word", "title_vi": "Bước khởi đầu", "desc_vi": "Thêm từ vựng đầu tiên", "unlocked": total >= 1},
        {"id": "quick_learner", "title_vi": "Học nhanh", "desc_vi": "Thuộc 5 từ vựng", "unlocked": learned >= 5},
        {"id": "word_master", "title_vi": "Bậc thầy từ vựng", "desc_vi": "Thuộc 50 từ vựng", "unlocked": learned >= 50},
        {"id": "streak_7", "title_vi": "Kiên trì", "desc_vi": "Duy trì chuỗi 7 ngày", "unlocked": streak >= 7},
        {"id": "centurion", "title_vi": "Trăm điểm", "desc_vi": "Đạt 100 điểm từ vựng", "unlocked": points >= 100}
    ]
    
    return {
        "total": total, "learned": learned, "completion": completion,
        "points": points, "streak": streak, "achievements": achievements
    }


def _writing_points(mode: str, score: float, num_items: int) -> int:
    """Points reward accuracy and effort. Sentence/paragraph: score(/10) x items.
    IELTS: band(/9) scaled up so a full essay is worth a meaningful chunk."""
    if mode == "ielts":
        return int(round(score * 20))          # band 7.0 -> 140 pts
    return int(round(score * max(1, num_items)))  # avg 8.0 over 10 -> 80 pts


@app.post("/api/writing/attempts")
def writing_log_attempt(req: WritingAttemptRequest, db=Depends(get_db)):
    points = _writing_points(req.mode, req.score, req.num_items)
    attempt = models.WritingAttempt(
        owner_email=req.owner_email, mode=req.mode, level=req.level,
        category=req.category, task=req.task, score=req.score,
        points=points, num_items=req.num_items,
    )
    db.add(attempt)
    db.commit()
    db.refresh(attempt)
    return {"id": attempt.id, "points_earned": points}


@app.get("/api/writing/progress")
def writing_progress(email: str | None = None, db=Depends(get_db)):
    q = db.query(models.WritingAttempt)
    if email:
        q = q.filter(models.WritingAttempt.owner_email == email)
    rows = q.order_by(models.WritingAttempt.created_at.desc()).all()

    total_points = sum(r.points or 0 for r in rows)
    sessions = len(rows)
    avg_score = round(sum(r.score or 0 for r in rows) / sessions, 1) if sessions else 0.0

    # Per-category accuracy (sentence/paragraph modes) — surfaces weak areas.
    by_cat = {}
    for r in rows:
        key = r.category or (r.task if r.mode == "ielts" else "Khác")
        if not key:
            continue
        acc = by_cat.setdefault(key, {"category": key, "sessions": 0, "avg_score": 0.0, "_sum": 0.0})
        acc["sessions"] += 1
        acc["_sum"] += (r.score or 0)
    categories = []
    for acc in by_cat.values():
        acc["avg_score"] = round(acc["_sum"] / acc["sessions"], 1) if acc["sessions"] else 0.0
        acc.pop("_sum", None)
        categories.append(acc)
    categories.sort(key=lambda c: c["avg_score"])   # weakest first

    # Streak: consecutive days (up to today) with at least one attempt.
    import datetime as _dt
    days = {r.created_at.date() for r in rows if r.created_at}
    streak = 0
    day = _dt.date.today()
    while day in days:
        streak += 1
        day -= _dt.timedelta(days=1)

    recent = [
        {"mode": r.mode, "level": r.level, "category": r.category, "task": r.task,
         "score": r.score, "points": r.points,
         "created_at": r.created_at.isoformat() if r.created_at else None}
        for r in rows[:8]
    ]

    return {
        "total_points": total_points, "sessions": sessions, "avg_score": avg_score,
        "streak": streak, "categories": categories, "recent": recent,
    }


@app.post("/api/teachback/start")
async def teachback_start(req: TeachBackStartRequest, db: Session = Depends(get_db)):
    from nlp import teachback as tb_engine
    
    context = req.context or ""
    if not context:
        if req.project_id:
            project = crud.get_project(db, req.project_id)
            if project:
                sources = [doc.content for doc in project.documents if doc.content]
                context = "\n\n---\n\n".join(sources)[:30000]
        elif req.document_id:
            doc = db.query(models.Document).filter(models.Document.id == req.document_id).first()
            if doc and doc.content:
                context = doc.content[:30000]

    data = None
    current_key = get_api_key(req.api_key)
    if current_key:
        try:
            prompt = (
                f"Bạn là giáo viên. Chọn MỘT chủ đề rõ ràng từ topic '{req.topic}' hoặc đoạn văn bản sau, "
                f"sau đó liệt kê 3-6 điểm chính (key_points) mà một lời giải thích tốt cần phải bao hàm. "
                f"Cuối cùng, hãy viết một câu hỏi (question_vi) bằng tiếng Việt yêu cầu người học giải thích lại chủ đề đó bằng lời của họ theo kỹ thuật Feynman.\n\n"
                f"Context: {context}"
            )
            config = LocalAgentConfig(
                api_key=current_key, model="gemini-3.1-flash-lite",
                response_schema=TeachBackPromptSchema, capabilities=FAST_CAPS,
            )
            result = await agent_run(config, prompt, structured=True, timeout=30)
            if result and result.get("concept"):
                data = result
        except Exception as e:
            print(f"[teachback start AI failed] {e}")

    if not data:
        data = tb_engine.make_prompt(req.topic, context, req.level or "")

    return data


@app.post("/api/teachback/evaluate")
async def teachback_evaluate(req: TeachBackEvalRequest):
    from nlp import teachback as tb_engine

    data = None
    current_key = get_api_key(req.api_key)
    if current_key:
        try:
            prompt = (
                f"Bạn là một gia sư áp dụng kỹ thuật Feynman. Hãy đánh giá lời giải thích của người học về khái niệm '{req.concept}'.\n"
                f"Các điểm cần có: {', '.join(req.key_points)}\n"
                f"Tài liệu tham khảo: {req.context or 'Không có'}\n"
                f"Lời giải thích của người học: {req.explanation}\n\n"
                "Trả về:\n"
                "- understanding_score: 0-100 (tỷ lệ hiểu/ghi nhớ)\n"
                "- covered: danh sách các điểm đã được giải thích tốt\n"
                "- gaps: danh sách các điểm quan trọng còn thiếu\n"
                "- misconceptions: danh sách các hiểu lầm (claim: người học nói gì, correction: sửa lại cho đúng)\n"
                "- followup_question: câu hỏi bằng tiếng Việt để đào sâu\n"
                "- feedback_vi: nhận xét chung, cụ thể, khích lệ (tiếng Việt)."
            )
            config = LocalAgentConfig(
                api_key=current_key, model="gemini-3.1-flash-lite",
                response_schema=TeachBackEvalSchema, capabilities=FAST_CAPS,
            )
            result = await agent_run(config, prompt, structured=True, timeout=30)
            if result and "understanding_score" in result:
                data = result
        except Exception as e:
            print(f"[teachback evaluate AI failed] {e}")

    if not data:
        data = tb_engine.evaluate(req.concept, req.key_points, req.explanation, req.context or "")

    return data


@app.post("/api/share")
def share_artifact(req: ShareRequest, db=Depends(get_db)):
    if req.artifact_id is None and req.drive_file_id is None:
        raise HTTPException(status_code=400, detail="Cần artifact_id hoặc drive_file_id")
    if req.artifact_id is not None:
        exists = db.query(models.Artifact).filter(models.Artifact.id == req.artifact_id).first()
        if not exists:
            raise HTTPException(status_code=404, detail="Artifact không tìm thấy")
    share = models.Share(
        artifact_id=req.artifact_id,
        drive_file_id=req.drive_file_id,
        drive_file_name=req.drive_file_name,
        owner_email=req.owner_email,
        shared_with_email=req.to_email,
    )
    db.add(share)
    db.commit()
    db.refresh(share)
    return {"success": True, "id": share.id}


@app.get("/api/studio/overview")
def studio_overview(email: str | None = None, db=Depends(get_db)):
    q = db.query(models.Artifact).filter(models.Artifact.type.in_(["exam", "examdoc"]))
    if email:
        q = q.filter(
            (models.Artifact.owner_email == email) | (models.Artifact.owner_email == None)
        )
    else:
        q = q.filter(models.Artifact.owner_email == None)
    mine_rows = q.order_by(models.Artifact.created_at.desc()).limit(20).all()
    mine = [
        {"id": a.id, "type": a.type, "title": a.title,
         "created_at": a.created_at.isoformat() if a.created_at else None}
        for a in mine_rows
    ]

    shared_exams, shared_examdocs, shared_library = [], [], []
    if email:
        shares = db.query(models.Share).filter(models.Share.shared_with_email == email).all()
        for s in shares:
            if s.drive_file_id and not s.artifact_id:
                shared_library.append({
                    "drive_file_id": s.drive_file_id,
                    "name": s.drive_file_name,
                    "from": s.owner_email,
                    "created_at": s.created_at.isoformat() if s.created_at else None,
                })
            elif s.artifact_id:
                art = db.query(models.Artifact).filter(models.Artifact.id == s.artifact_id).first()
                if not art:
                    continue
                entry = {
                    "id": art.id, "type": art.type, "title": art.title,
                    "created_at": art.created_at.isoformat() if art.created_at else None,
                    "from": s.owner_email,
                }
                if art.type == "exam":
                    shared_exams.append(entry)
                elif art.type == "examdoc":
                    shared_examdocs.append(entry)

    return {
        "mine": mine,
        "shared_exams": shared_exams,
        "shared_examdocs": shared_examdocs,
        "shared_library": shared_library,
    }


@app.post("/api/extract_text")
async def extract_text(file: UploadFile = File(...)):
    filename = file.filename or "upload"
    suffix = os.path.splitext(filename)[-1].lower()
    try:
        raw = await file.read()
        if suffix == ".pdf":
            import pdfplumber
            tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            try:
                tmp.write(raw)
                tmp.close()
                with pdfplumber.open(tmp.name) as pdf:
                    text = "\n".join(p.extract_text() or "" for p in pdf.pages)
            finally:
                os.unlink(tmp.name)
        elif suffix == ".docx":
            import docx
            tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            try:
                tmp.write(raw)
                tmp.close()
                d = docx.Document(tmp.name)
                text = "\n".join(p.text for p in d.paragraphs)
            finally:
                os.unlink(tmp.name)
        else:
            text = raw.decode("utf-8", errors="ignore")
        return {"filename": filename, "text": text, "chars": len(text)}
    except Exception as e:
        return {"filename": filename, "text": "", "error": str(e)}
# ── Team chat: human-to-human realtime-ish project group chat (drafted via 9Router) ──
from typing import Optional

class TeamMessageIn(BaseModel):
    author_email: Optional[str] = None
    author_name: str = "Ẩn danh"
    content: str

def _msg_dict(m):
    return {
        "id": m.id,
        "author_email": m.author_email,
        "author_name": m.author_name,
        "content": m.content,
        "created_at": m.created_at.isoformat() if m.created_at else None,
    }

@app.get("/api/projects/{project_id}/team_messages")
def get_team_messages(project_id: int, after_id: int = 0, db=Depends(get_db)):
    msgs = (
        db.query(models.TeamMessage)
        .filter(models.TeamMessage.project_id == project_id, models.TeamMessage.id > after_id)
        .order_by(models.TeamMessage.id.asc())
        .limit(200)
        .all()
    )
    return {"messages": [_msg_dict(m) for m in msgs]}

@app.post("/api/projects/{project_id}/team_messages")
def post_team_message(project_id: int, body: TeamMessageIn, db=Depends(get_db)):
    if not body.content or not body.content.strip():
        raise HTTPException(status_code=400, detail="content must not be empty")
    msg = models.TeamMessage(
        project_id=project_id,
        author_email=body.author_email,
        author_name=body.author_name,
        content=body.content.strip(),
        created_at=datetime.datetime.utcnow(),
    )
    db.add(msg)
    db.commit()
    db.refresh(msg)
    return _msg_dict(msg)